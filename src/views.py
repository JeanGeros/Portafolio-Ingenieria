from copyreg import constructor
import email
from multiprocessing.sharedctypes import Value
from this import d
from django.shortcuts import render, redirect
from django.http import HttpResponseRedirect
from django.contrib.auth import authenticate, login, logout
from django.views.decorators.csrf import csrf_exempt
from django.db import connection
import cx_Oracle

from datetime import datetime
import string
import json

import pandas as pd

from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth.models import User
import sweetify

import datetime
# import qrcode
from django.contrib.auth.decorators import login_required
from django.db.models import Sum

import openpyxl
from tempfile import NamedTemporaryFile
from datetime import datetime
import numpy as np
from django.http import HttpResponse
from django.http import FileResponse
import webbrowser
import os
from django.conf import settings
# from django.core.mail import EmailMultiAlternatives, send_mail
from django.template.loader import get_template
from django.template import loader
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from docx import Document
from django.http import HttpResponse

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4, A2, A3
from reportlab.lib.pagesizes import mm
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER 
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, NextPageTemplate

from reportlab.pdfbase.pdfmetrics import registerFont
from reportlab.pdfbase.ttfonts import TTFont
from  datetime import date
from reportlab.lib.units import inch
import reportlab
reportlab.rl_config.TTFSearchPath.append(str(settings.BASE_DIR) + '/src/lib/reportlabs/fonts')

from io import BytesIO
import re
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from django.contrib import messages

from src.static import images
from src.forms import (
    FormClienteNormal1, FormClienteNormal2, FormClienteNormal3, addproductsForm, FormVendedorPersona,
    FormVendedorUsuario, FormVendedorEmpleado, FormEmpleadoPersona, FormEmpleadoUsuario, FormEmpleadoEmpleado,
    FormProveedor, FormProductoproveedor, FormBodega, FormClienteEmpresa, FormCliente, FormTipodocumento, FormVenta, 
    FormDocu, FormProveedorUsuario,FormDireProvee, FormFamiliaProduct, FormDetalleorden, FormOrdencompra, FormAccionpagina
)

from .models import (
    Detalleorden, Estadoorden, Guiadespacho, Ordencompra, Persona, Direccion, Usuario, Cliente, Estado, Comuna, Tipobarrio, Despacho,
    Tipovivienda, Rolusuario, Direccioncliente, Empresa, Proveedor, Tipoproducto, Producto, Familiaproducto, Tipopago,
    Empleado, Cargo, Tiporubro, Recepcion, Productoproveedor, Bodega, Boleta, Factura, Venta, Tipodocumento, Detalleventa, Boleta,
    Notacredito, Accionpagina
)

from django.utils.encoding import smart_str

def Seguimiento_paginas(modulo, usuario):
    today = datetime.now()
    usuario_id = Usuario.objects.get(nombreusuario=usuario)
    seguimientoPag = Accionpagina.objects.create(
        fechain = today,
        modulo = modulo,
        usuarioid = usuario_id
    )

def Index(request):
    # Seguimiento_paginas("Inicio", request.user)
    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    context = {
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'index.html', context)

def Ingreso(request):
    # Seguimiento_paginas("Ingreso de usuarios", request.user)
    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    if request.method == 'POST':

        username = request.POST.get('usuario')
        password = request.POST.get('contrase??a')

        usuario_correo = Usuario.objects.filter(nombreusuario=username).exists()

        if usuario_correo == True:

            user = authenticate(request, username=username, password=password)

            if user is not None:

                login(request, user)
                sweetify.success(request, 'Ingreso realizado con exito')
                return redirect('index')

            else:

                sweetify.warning(request, 'Usuario y/o contrase??a inv??lidos.')
        else:

            sweetify.warning(request, 'Usuario y/o contrase??a inv??lidos.')

    context = {
        'tipo_usuario': tipo_usuario
    }
    
    return render(request, 'ingreso/ingreso_usuarios.html', context)

#************************************Productos*********************************************
@login_required(login_url="ingreso")
def Agregar_productos(request):
    Seguimiento_paginas("Modulo Productos - Crear Producto", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    form = addproductsForm(request.POST, request.FILES)
    form_prov = FormProductoproveedor(request.POST, request.FILES)
    form_bodega = FormBodega(request.POST, request.FILES)

    if request.method == 'POST':
        nombre = request.POST.get('nombre')
        precio = request.POST.get('precio')
        stock = request.POST.get('stock')
        stockcritico = request.POST.get('stockcritico')
        fechavencimiento = request.POST.get('fechavencimiento')
        imagen = request.FILES.get('imagen')
        proveedorid = request.POST.get('proveedorid')
        tipoproductoid = request.POST.get('tipoproductoid')
        familiaproid = request.POST.get('familiaproid')
        estadoid = request.POST.get('estadoid')
        bodega_id = request.POST.get('bodegaid')

        tipo_producto = Tipoproducto.objects.get(tipoproductoid=tipoproductoid)
        familia_producto = Familiaproducto.objects.get(familiaproid=familiaproid)
        estado_producto = Estado.objects.get(estadoid=estadoid)
        proveedor = Proveedor.objects.get(proveedorid=proveedorid)
        bodega = Bodega.objects.get(bodegaId=bodega_id)

        if len(fechavencimiento) == 10:
            fecha = fechavencimiento.split("/")
            fechavencimiento = f"{fecha[2]}-{fecha[1]}-{fecha[0]}"
            fechacodigo = f"{fecha[2]}{fecha[1]}{fecha[0]}"
        else:
            fechavencimiento = None
            fechacodigo = "00000000"

        codigo = f"{proveedor.proveedorid}{familia_producto.familiaproid}{fechacodigo}{tipo_producto.tipoproductoid}"
        try:
            product = Producto.objects.create(
                nombre=nombre.strip(),
                precio=precio.strip(),
                stock=stock.strip(),
                stockcritico=stockcritico.strip(),
                fechavencimiento=fechavencimiento,
                codigo=codigo,
                imagen=imagen,
                tipoproductoid=tipo_producto,
                familiaproid=familia_producto,
                estadoid=estado_producto,
                bodegaid=bodega
            )

            ultimo_producto = Producto.objects.order_by('productoid').last()
            prov_producto = Productoproveedor.objects.create(
                productoid=ultimo_producto,
                proveedorid=proveedor
            )

            if product is not None and prov_producto is not None:
                sweetify.success(request, 'Producto creado correctamente')
                return redirect('listar_productos')

        except Exception as error:
            sweetify.warning(request, error)

    context = {
        'form': form,
        'form2': form_prov,
        'form_bodega': form_bodega,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'productos/agregar_productos.html', context)


@login_required(login_url="ingreso")
def Listar_productos(request):

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    productos = Productoproveedor.objects.all()
    Seguimiento_paginas("Modulo Productos - Listar Productos", request.user)

    if request.method == 'POST':
        if request.POST.get('CambiarEstado') is not None:
            id_producto = request.POST.get('CambiarEstado')
            Cambiar_estado_producto(id_producto)
            producto = Producto.objects.get(productoid=id_producto)
            sweetify.success(request,
                             f'El producto {producto.nombre} ha quedado {producto.estadoid.descripcion} correctamente')
            return redirect('listar_productos')

        if request.POST.get('VerProducto') is not None:
            request.session['_old_post'] = request.POST
            return HttpResponseRedirect('ver_producto')

        if request.POST.get('EditarProducto') is not None:
            request.session['_old_post'] = request.POST
            return HttpResponseRedirect('editar_producto')

    context = {
        'productos': productos,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'productos/listar_productos.html', context)

def Cambiar_estado_producto(id_producto):
    
    producto = Producto.objects.get(productoid=id_producto)
    if producto.estadoid.descripcion == 'Activo':
        producto.estadoid = Estado.objects.get(descripcion="Inactivo")
        producto.save()
    else:
        producto.estadoid = Estado.objects.get(descripcion="Activo")
        producto.save()

@login_required(login_url="ingreso")
def Ver_producto(request):
    Seguimiento_paginas("Modulo Productos - Ver Detalles Producto", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    old_post = request.session.get('_old_post')
    producto = Productoproveedor.objects.get(productoid=old_post['VerProducto'])
    prov_producto = Productoproveedor.objects.get(productoid=old_post['VerProducto'],
                                                  proveedorid=old_post['proveedor'])
    context = {
        'producto': producto,
        'prov_producto': prov_producto,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'productos/ver_producto.html', context)


@login_required(login_url="ingreso")
def Editar_producto(request):
    Seguimiento_paginas("Modulo Productos - Editar Producto", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    old_post = request.session.get('_old_post')

    producto = Producto.objects.get(productoid=old_post['EditarProducto'])
    prov_producto = Productoproveedor.objects.get(productoid=old_post['EditarProducto'],
                                                  proveedorid=old_post['proveedor'])

    form2 = FormProductoproveedor(request.POST or None, instance=prov_producto)
    form = addproductsForm(request.POST or None, instance=producto)

    if request.method == 'POST':
        nombre = request.POST.get('nombre')
        precio = request.POST.get('precio')
        stock = request.POST.get('stock')
        stockcritico = request.POST.get('stockcritico')
        fechavencimiento = request.POST.get('fechavencimiento')
        imagen = request.FILES.get('imagen')
        proveedorid = request.POST.get('proveedorid')
        tipoproductoid = request.POST.get('tipoproductoid')
        familiaproid = request.POST.get('familiaproid')
        estadoid = request.POST.get('estadoid')
        bodega_id = request.POST.get('bodegaid')

        proveedor = Proveedor.objects.get(proveedorid=proveedorid)
        tipo_producto = Tipoproducto.objects.get(tipoproductoid=tipoproductoid)
        familia_producto = Familiaproducto.objects.get(familiaproid=familiaproid)
        Estado_producto = Estado.objects.get(estadoid=estadoid)
        bodega = Bodega.objects.get(bodegaId=bodega_id)

        fecha = fechavencimiento.split("/")
        if len(fechavencimiento) == 10:
            fechavencimiento = f"{fecha[2]}-{fecha[1]}-{fecha[0]}"
            fechacodigo = f"{fecha[2]}{fecha[1]}{fecha[0]}"
        else:
            fechavencimiento = None
            fechacodigo = "00000000"

        codigo = f"{proveedor.proveedorid}{familia_producto.familiaproid}{fechacodigo}{tipo_producto.tipoproductoid}"

        try:
            producto, created = Producto.objects.get_or_create(productoid=old_post['EditarProducto'])
            producto.nombre = nombre
            producto.precio = precio
            producto.stock = stock
            producto.stockcritico = stockcritico
            producto.fechavencimiento = fechavencimiento
            producto.tipoproductoid = tipo_producto
            producto.familiaproid = familia_producto
            producto.estadoid = Estado_producto
            producto.codigo = codigo
            producto.bodegaid = bodega

            if imagen:
                producto.imagen = imagen

            producto_prove, created = Productoproveedor.objects.get_or_create(productoid=old_post['EditarProducto'],
                                                                              proveedorid=old_post['proveedor'])
            producto_prove.proveedorid = proveedor

            producto_prove.save()
            producto.save()
            sweetify.success(request, 'Producto actualizado correctamente')
            return redirect('listar_productos')
        except Exception as error:
            sweetify.error(request, error)

    context = {
        'form': form,
        'form2': form2,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'productos/editar_productos.html', context)

##********************Perfiles*******************************************************************

@login_required(login_url="ingreso")
def Ver_perfil(request):
    Seguimiento_paginas("Ver Perfil", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    ver_perfil = request.session.get('_ver_perfil')
    datos_usuario = Usuario.objects.filter(nombreusuario=ver_perfil['VerPerfil']).values('nombreusuario','empresaid','personaid')
    usuario = Usuario.objects.get(nombreusuario=ver_perfil['VerPerfil'])
    persona = Persona.objects.get(personaid=datos_usuario[0]['personaid'])

    if Cliente.objects.filter(personaid=datos_usuario[0]['personaid']).exists():
        cliente = Cliente.objects.get(personaid=datos_usuario[0]['personaid'])
    else:
        cliente = None

    if datos_usuario[0]['empresaid'] is not None:
        empresa = Empresa.objects.get(empresaid=datos_usuario[0]['empresaid'])
    else:
        empresa = None

    if Direccioncliente.objects.filter(clienteid=cliente).exists():

        cliente = Cliente.objects.get(personaid=datos_usuario[0]['personaid'])
        direccion = Direccioncliente.objects.filter(clienteid=cliente).values('direccionid')
        perfil_direccion = []
        for x in direccion:
            perfil_direccion.append(Direccion.objects.get(direccionid=x['direccionid']))
        direccion = perfil_direccion
    else:
        direccion = None

    if request.method == 'POST':

        if request.POST.get('EditarPerfil') is not None:
            request.session['_editar_perfil'] = request.POST
            return redirect('editar_perfil')

        if request.POST.get('RevisarCompras') is not None:
            request.session['_revisar_compras'] = request.POST
            return redirect('revisar_compras')

        nombre_usuario = request.POST.get('BajarPerfil')
        
        usuario_desactivar = Usuario.objects.filter(nombreusuario=nombre_usuario).values('personaid')
        usuario_desactivar = usuario_desactivar[0]['personaid']
        estadoid = Estado.objects.get(descripcion="Inactivo")

        persona_in, created = Persona.objects.get_or_create(personaid=usuario_desactivar)
        persona_in.estadoid = estadoid
        persona_in.save()
        
        if Cliente.objects.filter(personaid=usuario_desactivar).exists():
            cliente_in, created = Cliente.objects.get_or_create(personaid=usuario_desactivar)
            cliente_in.estadoid = estadoid
            cliente_in.save()

        user_django = User.objects.get(username=nombre_usuario)
        user_django.is_active = False
        user_django.save()

        logout(request)
        return redirect('index')

    context = {
        'usuario': usuario,
        'persona': persona,
        'empresa': empresa,
        'cliente': cliente,
        'direccion': direccion,
        'tipo_usuario': tipo_usuario
    }

    return render(request, 'perfiles/ver_perfil.html', context)

@login_required(login_url="ingreso")
def Revisar_compras(request):
    Seguimiento_paginas("Revisar Compras", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    cliente= request.session.get('_revisar_compras')
    cliente = Persona.objects.filter(runcuerpo=cliente['run'], dv=cliente['dv']).values('personaid')
    cliente = Cliente.objects.get(personaid=cliente[0]['personaid'])

    lista = []
    lista.append(["Nro Venta","Fecha","Cant. productos","Total","Documento"])    
    val = Detalleventa.objects.filter(nroventa__clienteid=cliente).values('nroventa').annotate(suma=Sum('cantidad')).values_list('nroventa','nroventa__fechaventa','suma','nroventa__totalventa','nroventa__tipodocumentoid__descripcion')
    correo =  Usuario.objects.filter(nombreusuario=request.user).values('email')

    if request.method == 'POST':
        EnviarCorreo = request.POST.get('EnviarCorreo')

        if EnviarCorreo is not None:
            template = get_template('perfiles/descargar_compras.html')
            content = template.render({'val':val})
            
            message = loader.render_to_string(
                'perfiles/descargar_compras.html',
                {
                    'val': val
                }
            )
            de_email = settings.EMAIL_HOST_USER
            para_email = correo[0]['email']
            mime_message = MIMEText(message, "html", _charset="utf-8")
            mime_message["From"] = de_email
            mime_message["To"] = para_email
            mime_message["Subject"] = "Detalle de las compras realizadas"
            
            smtpObj = smtplib.SMTP(settings.EMAIL_HOST, 587)
            smtpObj.login(de_email, settings.EMAIL_HOST_PASSWORD)
            smtpObj.sendmail(de_email, para_email, mime_message.as_string())
            sweetify.success(request,"Las compras realizadas fueron enviadas a su correo")
            return redirect('revisar_compras')
        
        tipoInforme = request.POST.get('informeCheck')
        descargarInforme = request.POST.get('descargarInforme')
        
        if tipoInforme is not None and descargarInforme is not None:

            for valores in val:

                lista.append(list(valores))

            if tipoInforme == "informeExcel":

                nombre_archivo = "Compras"
                tipo_doc = 'ms-excel'
                extension = 'xlsx'
                
                return  creacion_excel(nombre_archivo, lista, tipo_doc, extension)

            if tipoInforme == "informePdf":
                
                tipo_doc = 'pdf'
                extension = 'pdf'
                nombre = 'Compras'
                
                return creacion_pdf(lista,tipo_doc,A4,nombre,extension, valor=True)

            if tipoInforme == "informeWord":

                nombre_archivo = "Compras"
                return  creacion_doc(lista, nombre_archivo)
    
    if val.exists() == False:
        val = None

    context = {
        'tipo_usuario': tipo_usuario,
        'val': val
    }

    return render(request, 'perfiles/revisar_compras.html', context)

@login_required(login_url="ingreso")
def Editar_perfil(request):
    Seguimiento_paginas("Editar Perfil", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    editar_perfil= request.session.get('_editar_perfil')

    persona = Persona.objects.filter(runcuerpo=editar_perfil['run'],dv=editar_perfil['dv']).values('personaid')
    usuario = Usuario.objects.filter(personaid=persona[0]['personaid']).values('empresaid','email')
    if Cliente.objects.filter(personaid=persona[0]['personaid']).exists():
        cliente = Cliente.objects.get(personaid=persona[0]['personaid'])
        direccion = Direccioncliente.objects.filter(clienteid=cliente).values('direccionid')
        perfil_direccion = []
        for x in direccion:
            perfil_direccion.append(Direccion.objects.get(direccionid=x['direccionid']))
        form2 = []
    else: 
        perfil_direccion = None

    perfil_usuario = Usuario.objects.get(personaid=persona[0]['personaid'])
    perfil_persona = Persona.objects.get(personaid=persona[0]['personaid'])

    if Empresa.objects.filter(empresaid=usuario[0]['empresaid']).exists():
        perfil_empresa = Empresa.objects.get(empresaid=usuario[0]['empresaid'])
    else:
        perfil_empresa = None

    form1 = FormClienteNormal1(request.POST or None, instance=perfil_persona)

    if perfil_direccion is not None:
        for x in perfil_direccion:
            form2.append(FormClienteNormal2(request.POST or None, instance=x))
    else: 
        form2 = None

    form3 = FormClienteNormal3(request.POST or None, instance=perfil_usuario)

    if perfil_empresa is not None:
        form4  = FormClienteEmpresa(request.POST or None, instance=perfil_empresa)
    else:
        form4 = None

    if request.method == 'POST':

        run_cuerpo = request.POST.get('runcuerpo')
        dv = request.POST.get('dv')
        nombres = request.POST.get('nombres')
        apellido_paterno = request.POST.get('apellidopaterno')
        apellido_materno = request.POST.get('apellidomaterno')
        telefono = request.POST.get('telefono')
        email = request.POST.get('email')

        if perfil_direccion is not None:
            calle = request.POST.get('calle')
            numero = request.POST.get('numero')
            comuna_id = request.POST.get('comunaid')
            tipo_vivienda_id = request.POST.get('tipoviviendaid')
            tipo_barrio_id = request.POST.get('tipobarrioid')
            nombre_sector = request.POST.get('nombresector')
            # nombre_usuario = request.POST.get('nombreusuario')

        if perfil_empresa is not None:
            razon_social = request.POST.get('razonsocial')
            rut_cuerpo = request.POST.get('rutcuerpo')
            dv = request.POST.get('dv')
            fono = request.POST.get('fono')

        perfil_persona, created = Persona.objects.get_or_create(personaid=persona[0]['personaid'])
        perfil_persona.runcuerpo = run_cuerpo
        perfil_persona.dv = dv
        perfil_persona.apellidopaterno = apellido_paterno
        perfil_persona.apellidomaterno = apellido_materno
        perfil_persona.nombres = nombres
        perfil_persona.telefono = telefono
        perfil_persona.save()

        if perfil_direccion is not None:

            perfil_direccion, created = Direccion.objects.get_or_create(direccionid=direccion[0]['direccionid'])
            perfil_direccion.calle = calle
            perfil_direccion.numero = numero
            perfil_direccion.comunaid = Comuna.objects.get(comunaid=comuna_id)
            perfil_direccion.tipoviviendaid = Tipovivienda.objects.get(tipoviviendaid=tipo_vivienda_id)
            perfil_direccion.tipobarrioid = Tipobarrio.objects.get(tipobarrioid=tipo_barrio_id)
            perfil_direccion.nombresector = nombre_sector
            perfil_direccion.save()

        if perfil_empresa is not None:
            
            perfil_empresa, created = Empresa.objects.get_or_create(empresaid=usuario[0]['empresaid'])
            perfil_empresa.razonsocial = razon_social
            perfil_empresa.rutcuerpo = rut_cuerpo
            perfil_empresa.dv = dv
            perfil_empresa.fono = fono
            perfil_empresa.save()

        perfil_usuario, created = Usuario.objects.get_or_create(personaid=persona[0]['personaid'])
        perfil_usuario.email = email
        perfil_usuario.save()
        
        user_django = User.objects.get(email=usuario[0]['email'])
        user_django.username = user_django.username
        user_django.first_name = nombres
        user_django.last_name = apellido_paterno
        user_django.email = email
        user_django.save()

        sweetify.success(request,"Perfil actualizado correctamente")
        return redirect('ver_perfil')
    
    context = {
        'form1': form1,
        'form2': form2,
        'form3': form3,
        'form4': form4,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'perfiles/editar_perfil.html', context)

##********************Clientes*******************************************************************

def Seleccion_registro(request):

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None
    
    context = {
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'clientes/seleccion_registro.html', context)


def Registro_clientes(request):
    # Seguimiento_paginas("Modulo Clientes - Registro Clientes", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    form1 = FormClienteNormal1()
    form2 = FormClienteNormal2()
    form3 = FormClienteNormal3()

    if request.method == 'POST':
        run_cuerpo = request.POST.get('runcuerpo')
        dv = request.POST.get('dv')
        nombres = request.POST.get('nombres')
        apellido_paterno = request.POST.get('apellidopaterno')
        apellido_materno = request.POST.get('apellidomaterno')
        telefono = request.POST.get('telefono')
        nombre_usuario = request.POST.get('nombreusuario')
        email = request.POST.get('email')
        calle = request.POST.get('calle')
        numero = request.POST.get('numero')
        comuna_id = request.POST.get('comunaid')
        tipo_vivienda_id = request.POST.get('tipoviviendaid')
        tipo_barrio_id = request.POST.get('tipobarrioid')
        nombre_sector = request.POST.get('nombresector')
        contrase??a = request.POST.get('password')
        confirme_contrase??a = request.POST.get('confirme_contrase??a')

        personaRegistro = Persona.objects.filter(runcuerpo=run_cuerpo, dv=dv).exists()
        usuarioRegistro = Usuario.objects.filter(email=email).exists()
        usuarioRegistro3 = Usuario.objects.filter(nombreusuario=nombre_usuario).exists()

        if personaRegistro == True:
            sweetify.warning(request, "El run ingresado ya existe")
        elif usuarioRegistro == True:
            sweetify.warning(request, "El correo ingresado ya existe")
        elif usuarioRegistro3 == True:
            sweetify.warning(request, "El usuario ingresado ya esta registrado")
        else:

            rolusuario = Rolusuario.objects.filter(descripcion="Cliente").values('rolid')

            django_cursor = connection.cursor()
            cursor = django_cursor.connection.cursor()
            procedimiento = cursor.callproc(
                'SP_InsertClientePersona', 
                [
                    int(run_cuerpo), 
                    dv, 
                    apellido_paterno, 
                    apellido_materno, 
                    nombres, 
                    int(telefono), 
                    rolusuario[0]['rolid'], 
                    calle,
                    str(numero),
                    nombre_sector,
                    tipo_vivienda_id,
                    tipo_barrio_id,
                    comuna_id,
                    email,
                    nombre_usuario,
                    contrase??a
                ]
            )

            user = User.objects.create_user(
                username=nombre_usuario,
                first_name=nombres,
                last_name=apellido_paterno,
                email=email,
                is_superuser=False,
                is_active=True
            )
            user.set_password(contrase??a)
            user.set_password(confirme_contrase??a)

            if (procedimiento is not None and 
                user is not None):
                user.save()

                sweetify.success(request, "Se ha registrado correctamente")
                return redirect('index')
            else:
                sweetify.error(request, "No es posible registrarse en este momento")

    context = {
        'form1': form1,
        'form2': form2,
        'form3': form3,
        'tipo_usuario': tipo_usuario
    }

    return render(request, 'clientes/registro_clientes.html', context)

@login_required(login_url="ingreso")
def Agregar_cliente(request):
    Seguimiento_paginas("Modulo Clientes - Agregar Clientes", request.user)
    
    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')
    
    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    form1 = FormClienteNormal1()
    form2 = FormClienteNormal2()
    form3 = FormClienteNormal3()

    if request.method == 'POST':

        run_cuerpo = request.POST.get('runcuerpo')
        dv = request.POST.get('dv')
        nombres = request.POST.get('nombres')
        apellido_paterno = request.POST.get('apellidopaterno')
        apellido_materno = request.POST.get('apellidomaterno')
        telefono = request.POST.get('telefono')
        nombre_usuario = request.POST.get('nombreusuario')
        email = request.POST.get('email')
        calle = request.POST.get('calle')
        numero = request.POST.get('numero')
        comuna_id = request.POST.get('comunaid')
        tipo_vivienda_id = request.POST.get('tipoviviendaid')
        tipo_barrio_id = request.POST.get('tipobarrioid')
        nombre_sector = request.POST.get('nombresector')
        contrase??a = request.POST.get('password')
        confirme_contrase??a = request.POST.get('confirme_contrase??a')

        personaRegistro = Persona.objects.filter(runcuerpo=run_cuerpo, dv=dv).exists()
        usuarioRegistro = Usuario.objects.filter(email=email).exists()
        usuarioRegistro3 = Usuario.objects.filter(nombreusuario=nombre_usuario).exists()

        if personaRegistro == True:
            sweetify.warning(request, "El run ingresado ya existe")
        elif usuarioRegistro == True:
            sweetify.warning(request, "El correo ingresado ya existe")
        elif usuarioRegistro3 == True:
            sweetify.warning(request, "El cliente ya cuenta con un usuario")
        else: 
            rolusuario = Rolusuario.objects.filter(descripcion="Cliente").values('rolid')

            django_cursor = connection.cursor()
            cursor = django_cursor.connection.cursor()
            procedimiento = cursor.callproc(
                'SP_InsertClientePersona', 
                [
                    int(run_cuerpo), 
                    dv, 
                    apellido_paterno, 
                    apellido_materno, 
                    nombres, 
                    int(telefono), 
                    rolusuario[0]['rolid'], 
                    calle,
                    str(numero),
                    nombre_sector,
                    tipo_vivienda_id,
                    tipo_barrio_id,
                    comuna_id,
                    email,
                    nombre_usuario,
                    contrase??a
                ]
            )

            user = User.objects.create_user(
                username=nombre_usuario,
                first_name=nombres,
                last_name=apellido_paterno,
                email=email,
                is_superuser=False,
                is_active=True
            )
            user.set_password(contrase??a)
            user.set_password(confirme_contrase??a)

            if (procedimiento is not None and 
                user is not None):
                user.save()

                sweetify.success(request, "Cliente creado correctamente")
                return redirect('listar_clientes')
            else:
                sweetify.error(request, "No es posible crear el cliente en este momento")

    context = {
        'form1': form1,
        'form2': form2,
        'form3': form3,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'clientes/agregar_cliente.html', context)

@login_required(login_url="ingreso")
def Listar_clientes(request):
    Seguimiento_paginas("Modulo Clientes - Lista de Clientes", request.user)
    
    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    clientes = Cliente.objects.all()

    if request.method == 'POST':

        if request.POST.get('CambiarEstado') is not None:
            id_cliente = request.POST.get('CambiarEstado')
            Cambiar_estado_cliente(id_cliente)
            cliente = Cliente.objects.get(clienteid=id_cliente)
            sweetify.success(request,
                             f'El cliente {cliente.personaid.runcuerpo} - {cliente.personaid.dv} ha quedado {cliente.estadoid.descripcion} correctamente')
            return redirect('listar_clientes')

        if request.POST.get('VerCliente') is not None:
            request.session['_old_post'] = request.POST
            return HttpResponseRedirect('ver_cliente')

        if request.POST.get('EditarCliente') is not None:
            request.session['_old_post'] = request.POST
            return HttpResponseRedirect('editar_cliente')

    context = {
        'clientes':clientes,
        'tipo_usuario': tipo_usuario
    }

    return render(request, 'clientes/listar_clientes.html', context)

def Cambiar_estado_cliente(id_cliente):
    cliente = Cliente.objects.get(clienteid=id_cliente)
    persona_id = Cliente.objects.filter(clienteid=id_cliente).values('personaid')
    usuario = Usuario.objects.filter(personaid=persona_id[0]['personaid']).values('nombreusuario')
    user = User.objects.get(username=usuario[0]['nombreusuario'])

    if cliente.estadoid.descripcion == 'Activo':
        cliente.estadoid = Estado.objects.get(descripcion="Inactivo")
        cliente.save()
        user.is_active = False
        user.save()
    else:
        cliente.estadoid = Estado.objects.get(descripcion="Activo")
        cliente.save()
        user.is_active = True
        user.save()

@login_required(login_url="ingreso")
def Ver_cliente(request):
    Seguimiento_paginas("Modulo Clientes - Ver Cliente", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    old_post = request.session.get('_old_post')
    ver_cliente_post = Cliente.objects.get(clienteid=old_post['VerCliente'])
    direccion = Direccioncliente.objects.filter(clienteid=ver_cliente_post).values('direccionid')
    direccion = Direccion.objects.get(direccionid=direccion[0]['direccionid'])
    cliente = Cliente.objects.get(clienteid=old_post['VerCliente'])

    context = {
        'tipo_usuario': tipo_usuario,
        'cliente': cliente,
        'direccion': direccion
    }

    return render(request, 'clientes/ver_cliente.html', context)

@login_required(login_url="ingreso")
def Editar_cliente(request):
    Seguimiento_paginas("Modulo Clientes - Editar Cliente", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    old_post = request.session.get('_old_post')

    cliente = Cliente.objects.filter(clienteid=old_post['EditarCliente']).values('personaid')
    cliente_persona = Persona.objects.get(personaid=cliente[0]['personaid'])
    cliente_usuario = Usuario.objects.get(personaid=cliente[0]['personaid'])
    direccion1 = Direccioncliente.objects.filter(clienteid=old_post['EditarCliente']).values('direccionid')
    direccion = Direccion.objects.get(direccionid=direccion1[0]['direccionid'])
    email_cliente = Usuario.objects.filter(personaid=cliente[0]['personaid']).values('email')

    form1 = FormClienteNormal1(request.POST or None, instance=cliente_persona)
    form2 = FormClienteNormal2(request.POST or None, instance=direccion)
    form3 = FormClienteNormal3(request.POST or None, instance=cliente_usuario)

    if request.method == 'POST':

        run_cuerpo = request.POST.get('runcuerpo')
        dv = request.POST.get('dv')
        nombres = request.POST.get('nombres')
        apellido_paterno = request.POST.get('apellidopaterno')
        apellido_materno = request.POST.get('apellidomaterno')
        telefono = request.POST.get('telefono')
        email = request.POST.get('email')
        calle = request.POST.get('calle')
        numero = request.POST.get('numero')
        comuna_id = request.POST.get('comunaid')
        tipo_vivienda_id = request.POST.get('tipoviviendaid')
        tipo_barrio_id = request.POST.get('tipobarrioid')
        nombre_sector = request.POST.get('nombresector')
        # nombre_usuario = request.POST.get('nombreusuario')

        cliente_persona, created = Persona.objects.get_or_create(personaid=cliente[0]['personaid'])
        cliente_persona.runcuerpo = run_cuerpo
        cliente_persona.dv = dv
        cliente_persona.apellidopaterno = apellido_paterno
        cliente_persona.apellidomaterno = apellido_materno
        cliente_persona.nombres = nombres
        cliente_persona.telefono = telefono
        cliente_persona.save()

        cliente_direccion, created = Direccion.objects.get_or_create(direccionid=direccion1[0]['direccionid'])
        cliente_direccion.calle = calle
        cliente_direccion.numero = numero
        cliente_direccion.comunaid = Comuna.objects.get(comunaid=comuna_id)
        cliente_direccion.tipoviviendaid = Tipovivienda.objects.get(tipoviviendaid=tipo_vivienda_id)
        cliente_direccion.tipobarrioid = Tipobarrio.objects.get(tipobarrioid=tipo_barrio_id)
        cliente_direccion.nombresector = nombre_sector
        cliente_direccion.save()

        cliente_usuario, created = Usuario.objects.get_or_create(personaid=cliente[0]['personaid'])
        cliente_usuario.email = email
        cliente_usuario.save()

        user_django = User.objects.get(email=email_cliente[0]['email'])
        user_django.username = user_django.username
        user_django.first_name = nombres
        user_django.last_name = apellido_paterno
        user_django.email = email
        user_django.save()

        sweetify.success(request, "Cliente actualizado correctamente")
        return redirect('listar_clientes')

    context = {
        'form1': form1,
        'form2': form2,
        'form3': form3,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'clientes/editar_cliente.html', context)

def Registro_clientes_empresa(request):

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None
    
    form2 = FormClienteNormal2()
    form3 = FormClienteNormal3()
    form4 = FormClienteEmpresa()

    if request.method == 'POST':
        
        dv = request.POST.get('dv')
        rut_cuerpo = request.POST.get('rutcuerpo')
        nombre_usuario = request.POST.get('nombreusuario')
        email = request.POST.get('email')
        calle = request.POST.get('calle')
        numero = request.POST.get('numero')
        comuna_id = request.POST.get('comunaid')
        tipo_vivienda_id = request.POST.get('tipoviviendaid')
        tipo_barrio_id = request.POST.get('tipobarrioid')
        nombre_sector = request.POST.get('nombresector')
        contrase??a = request.POST.get('password')
        confirme_contrase??a = request.POST.get('confirme_contrase??a')
        razon_social = request.POST.get('razonsocial')
        fono = request.POST.get('fono')

        usuarioRegistro = Usuario.objects.filter(email = email).exists()
        usuarioRegistro3 = Usuario.objects.filter(nombreusuario = nombre_usuario).exists()
        usuarioRegistroEmp = Empresa.objects.filter(razonsocial = razon_social).exists()

        if usuarioRegistro == True:
            sweetify.warning(request,"El correo ingresado ya existe")
        elif usuarioRegistro3 == True:
            sweetify.warning(request,"El usuario ingresado ya esta registrado")
        elif usuarioRegistroEmp == True:
            sweetify.warning(request,"La raz??n social ingresada ya existe")
        else: 
            rolusuario = Rolusuario.objects.filter(descripcion="Cliente").values('rolid')

            django_cursor = connection.cursor()
            cursor = django_cursor.connection.cursor()
            procedimiento = cursor.callproc(
                'SP_InsertClienteEmpresa', 
                [
                    razon_social,
                    int(rut_cuerpo), 
                    dv, 
                    fono,
                    rolusuario[0]['rolid'], 
                    calle,
                    str(numero),
                    nombre_sector,
                    tipo_vivienda_id,
                    tipo_barrio_id,
                    comuna_id,
                    email,
                    nombre_usuario,
                    contrase??a
                ]
            )

            user = User.objects.create_user(
                username = nombre_usuario,
                email = email,
                is_superuser = False,
                is_active = True
            )
            user.set_password(contrase??a)
            user.set_password(confirme_contrase??a)
            
            if (procedimiento is not None and 
                user is not None):
                
                user.save()

                sweetify.success(request,"Se ha registrado correctamente")
                return redirect('index')
            else:
                sweetify.error(request,"No es posible registrarse en este momento")

    context = {
        'form2': form2,
        'form3': form3,
        'form4': form4,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'clientes/registro_clientes_empresa.html', context)

##******************************************************************************************

##********************************Vendedores************************************************

@login_required(login_url="ingreso")
def Listar_vendedores(request):
    Seguimiento_paginas("Modulo Vendedores - Listar Vendedores", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    vendedores = Cargo.objects.get(descripcion="Vendedor")
    vendedores = Empleado.objects.filter(cargoid=vendedores)

    if request.method == 'POST':

        if request.POST.get('CambiarEstado') is not None:
            id_vendedor = request.POST.get('CambiarEstado')
            Cambiar_estado_vendedor(id_vendedor)
            vendedor = Empleado.objects.get(empleadoid=id_vendedor)
            sweetify.success(request,
                             f'El vendedor {vendedor.personaid.runcuerpo} - {vendedor.personaid.dv} ha quedado {vendedor.estadoid.descripcion} correctamente')
            return redirect('listar_vendedores')

        if request.POST.get('VerVendedor') is not None:
            request.session['_old_post'] = request.POST
            return HttpResponseRedirect('ver_vendedor')

        if request.POST.get('EditarVendedor') is not None:
            request.session['_old_post'] = request.POST
            return HttpResponseRedirect('editar_vendedor')

    context = {
        'vendedores':vendedores,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'vendedores/listar_vendedores.html', context)

@login_required(login_url="ingreso")
def Agregar_vendedor(request):
    Seguimiento_paginas("Modulo Vendedores - Agregar Vendedor", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    form1 = FormVendedorPersona()
    form2 = FormVendedorUsuario()
    form3 = FormVendedorEmpleado()

    if request.method == 'POST':
        run_cuerpo = request.POST.get('runcuerpo')
        dv = request.POST.get('dv')
        nombres = request.POST.get('nombres')
        apellido_paterno = request.POST.get('apellidopaterno')
        apellido_materno = request.POST.get('apellidomaterno')
        telefono = request.POST.get('telefono')
        nombre_usuario = request.POST.get('nombreusuario')
        fecha_ingreso = request.POST.get('fechaingreso')
        email = request.POST.get('email')
        contrase??a = request.POST.get('password')
        confirme_contrase??a = request.POST.get('confirme_contrase??a')

        fecha_ingreso = fecha_ingreso[6:10] + '-' + fecha_ingreso[3:5] + '-' + fecha_ingreso[0:2]

        personaRegistro = Persona.objects.filter(runcuerpo=run_cuerpo, dv=dv).exists()
        usuarioRegistro = Usuario.objects.filter(email=email).exists()
        usuarioRegistro3 = Usuario.objects.filter(nombreusuario=nombre_usuario).exists()

        if personaRegistro:
            sweetify.error(request, "El run ingresado ya existe")

        elif usuarioRegistro:
            sweetify.error(request, "El correo ingresado ya existe")

        elif usuarioRegistro3:
            sweetify.error(request, "El vendedor ya cuenta con un usuario")

        else:
            cargo = Cargo.objects.filter(descripcion="Vendedor").values('cargoid')
            rolusuario = Rolusuario.objects.filter(descripcion="Vendedor").values('rolid')

            django_cursor = connection.cursor()
            cursor = django_cursor.connection.cursor()
            procedimiento = cursor.callproc(
                'SP_Insertempleado', 
                [
                    int(run_cuerpo), 
                    dv, 
                    apellido_paterno, 
                    apellido_materno, 
                    nombres, 
                    int(telefono), 
                    fecha_ingreso, 
                    int(cargo[0]['cargoid']), 
                    email, 
                    nombre_usuario, 
                    contrase??a, 
                    rolusuario[0]['rolid']
                ]
            )

            user = User.objects.create_user(
                username=nombre_usuario,
                first_name=nombres,
                last_name=apellido_paterno,
                email=email,
                is_superuser=False,
                is_active=True
            )
            user.set_password(contrase??a)
            user.set_password(confirme_contrase??a)

            if (procedimiento is not None and 
                user is not None):
                user.save()
                sweetify.success(request, "Vendedor creado correctamente")
                return redirect('listar_vendedores')

            else:
                sweetify.error(request, "No es posible crear el vendedor en este momento")

    context = {
        'form1': form1,
        'form2': form2,
        'form3': form3,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'vendedores/agregar_vendedor.html', context)

@login_required(login_url="ingreso")
def Ver_vendedor(request):
    Seguimiento_paginas("Modulo Vendedores - Ver Vendedores", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    old_post = request.session.get('_old_post')
    vendedor = Empleado.objects.get(empleadoid=old_post['VerVendedor'])

    context = {
        'vendedor':vendedor,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'vendedores/ver_vendedor.html', context)

@login_required(login_url="ingreso")
def Editar_vendedor(request):
    Seguimiento_paginas("Modulo Vendedores - Editar Vendedor", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    old_post = request.session.get('_old_post')

    vendedor = Empleado.objects.filter(empleadoid=old_post['EditarVendedor']).values('personaid')
    vendedor_usuario = Usuario.objects.get(personaid=vendedor[0]['personaid'])
    vendedor_persona = Persona.objects.get(personaid=vendedor[0]['personaid'])
    vendedor_empleado = Empleado.objects.get(personaid=vendedor[0]['personaid'])

    form1 = FormVendedorPersona(request.POST or None, instance=vendedor_persona)
    form2 = FormVendedorUsuario(request.POST or None, instance=vendedor_usuario)
    form3 = FormVendedorEmpleado(request.POST or None, instance=vendedor_empleado)

    if request.method == 'POST':
        run_cuerpo = request.POST.get('runcuerpo')
        dv = request.POST.get('dv')
        nombres = request.POST.get('nombres')
        apellido_paterno = request.POST.get('apellidopaterno')
        apellido_materno = request.POST.get('apellidomaterno')
        telefono = request.POST.get('telefono')
        nombre_usuario = request.POST.get('nombreusuario')
        fecha_ingreso = request.POST.get('fechaingreso')
        email = request.POST.get('email')

        fecha_ingreso = fecha_ingreso[6:10] + '-' + fecha_ingreso[3:5] + '-' + fecha_ingreso[0:2]

        empleado, created = Empleado.objects.get_or_create(empleadoid=old_post['EditarVendedor'])
        empleado.fechaingreso = fecha_ingreso
        empleado.personaid.runcuerpo = run_cuerpo
        empleado.personaid.dv = dv
        empleado.personaid.nombres = nombres
        empleado.personaid.apellidopaterno = apellido_paterno
        empleado.personaid.apellidomaterno = apellido_materno
        empleado.personaid.telefono = telefono
        empleado.save()

        empleado_para_persona = Empleado.objects.filter(empleadoid=old_post['EditarVendedor']).values('personaid')
        usuario, created = Usuario.objects.get_or_create(personaid=empleado_para_persona[0]['personaid'])
        usuario.nombreusuario = nombre_usuario
        usuario.email = email
        usuario.save()

        persona, created = Persona.objects.get_or_create(personaid=empleado_para_persona[0]['personaid'])
        persona.runcuerpo = run_cuerpo
        persona.dv = dv
        persona.apellidopaterno = apellido_paterno
        persona.apellidomaterno = apellido_materno
        persona.nombres = nombres
        persona.telefono = telefono
        persona.save()

        user_django = User.objects.get(email=vendedor_usuario)
        user_django.username = user_django.username
        user_django.first_name = nombres
        user_django.last_name = apellido_paterno
        user_django.email = email
        user_django.save()

        sweetify.success(request, "Vendedor actualizado correctamente")
        return redirect('listar_vendedores')

    context = {
        'form1': form1,
        'form2': form2,
        'form3': form3,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'vendedores/editar_vendedor.html', context)

def Cambiar_estado_vendedor(id_vendedor):
    vendedor = Empleado.objects.get(empleadoid=id_vendedor)
    persona_id = Empleado.objects.filter(empleadoid=id_vendedor).values('personaid')
    usuario = Usuario.objects.filter(personaid=persona_id[0]['personaid']).values('nombreusuario')
    user = User.objects.get(username=usuario[0]['nombreusuario'])

    if vendedor.estadoid.descripcion == 'Activo':
        vendedor.estadoid = Estado.objects.get(descripcion="Inactivo")
        vendedor.save()
        user.is_active = False
        user.save()
    else:
        vendedor.estadoid = Estado.objects.get(descripcion="Activo")
        vendedor.save()
        user.is_active = True
        user.save()

@login_required(login_url="ingreso")
def Listar_clientes_vendedor(request):
    Seguimiento_paginas("Modulo Vendedores - Listar Cliente de Vendedor", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    clientes = Cliente.objects.all()

    if request.method == 'POST':

        if request.POST.get('VerCliente') is not None:
            request.session['_ver_cliente'] = request.POST
            return HttpResponseRedirect('ver_cliente_vendedor')

    context = {
        'tipo_usuario': tipo_usuario,
        'clientes': clientes,
    }

    return render(request, 'vendedores/listar_clientes_vendedor.html', context)

@login_required(login_url="ingreso")
def Ver_cliente_vendedor(request):
    Seguimiento_paginas("Modulo Vendedores - Ver Cliente Vendedor", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    ver_cliente_post = request.session.get('_ver_cliente')

    ver_cliente = Persona.objects.filter(runcuerpo=ver_cliente_post['run'], dv=ver_cliente_post['dv']).values('personaid')
    cliente_id = Cliente.objects.filter(personaid=ver_cliente[0]['personaid']).values('clienteid')
    direccion = Direccioncliente.objects.filter(clienteid=cliente_id[0]['clienteid']).values('direccionid')
    direccion = Direccion.objects.get(direccionid=direccion[0]['direccionid'])
    cliente = Cliente.objects.get(personaid=ver_cliente[0]['personaid'])

    if request.method == 'POST':

        if request.POST.get('EditarCliente') is not None:
            request.session['_editar_cliente'] = request.POST
            return HttpResponseRedirect('editar_cliente_vendedor')

    context = {
        'tipo_usuario': tipo_usuario,
        'cliente': cliente,
        'direccion': direccion
    }

    return render(request, 'vendedores/ver_cliente_vendedor.html', context)

@login_required(login_url="ingreso")
def Agregar_cliente_vendedor(request):
    Seguimiento_paginas("Modulo Vendedores - Agregar Cliente a Vendedor", request.user)
    
    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')
    
    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    form1 = FormClienteNormal1()
    form2 = FormClienteNormal2()
    form3 = FormClienteNormal3()

    if request.method == 'POST':
        
        run_cuerpo = request.POST.get('runcuerpo')
        dv = request.POST.get('dv')
        nombres = request.POST.get('nombres')
        apellido_paterno = request.POST.get('apellidopaterno')
        apellido_materno = request.POST.get('apellidomaterno')
        telefono = request.POST.get('telefono')
        nombre_usuario = request.POST.get('nombreusuario')
        email = request.POST.get('email')
        calle = request.POST.get('calle')
        numero = request.POST.get('numero')
        comuna_id = request.POST.get('comunaid')
        tipo_vivienda_id = request.POST.get('tipoviviendaid')
        tipo_barrio_id = request.POST.get('tipobarrioid')
        nombre_sector = request.POST.get('nombresector')
        contrase??a = request.POST.get('password')
        confirme_contrase??a = request.POST.get('confirme_contrase??a')

        personaRegistro = Persona.objects.filter(runcuerpo=run_cuerpo, dv=dv).exists()
        usuarioRegistro = Usuario.objects.filter(email = email).exists()
        usuarioRegistro3 = Usuario.objects.filter(nombreusuario = nombre_usuario).exists()

        if personaRegistro == True:
            sweetify.warning(request,"El run ingresado ya existe")
        elif usuarioRegistro == True:
            sweetify.warning(request,"El correo ingresado ya existe")
        elif usuarioRegistro3 == True:
            sweetify.warning(request,"El cliente ya cuenta con un usuario")
        else: 
            rolusuario = Rolusuario.objects.filter(descripcion="Cliente").values('rolid')

            django_cursor = connection.cursor()
            cursor = django_cursor.connection.cursor()
            procedimiento = cursor.callproc(
                'SP_InsertClientePersona', 
                [
                    int(run_cuerpo), 
                    dv, 
                    apellido_paterno, 
                    apellido_materno, 
                    nombres, 
                    int(telefono), 
                    rolusuario[0]['rolid'], 
                    calle,
                    str(numero),
                    nombre_sector,
                    tipo_vivienda_id,
                    tipo_barrio_id,
                    comuna_id,
                    email,
                    nombre_usuario,
                    contrase??a
                ]
            )

            user = User.objects.create_user(
                username=nombre_usuario,
                first_name=nombres,
                last_name=apellido_paterno,
                email=email,
                is_superuser=False,
                is_active=True
            )
            user.set_password(contrase??a)
            user.set_password(confirme_contrase??a)

            if (procedimiento is not None and 
                user is not None):
                user.save()

                sweetify.success(request,"Cliente creado correctamente")
                return redirect('registro_clientes')
            else:
                sweetify.error(request,"No es posible crear el cliente en este momento")

    context = {
        'form1': form1,
        'form2': form2,
        'form3': form3,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'vendedores/agregar_cliente_vendedor.html', context)

@login_required(login_url="ingreso")
def Editar_cliente_vendedor(request):
    Seguimiento_paginas("Modulo Vendedores - Editar Cliente de Vendedor", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    editar_cliente = request.session.get('_editar_cliente')
    
    cliente = Cliente.objects.filter(clienteid=editar_cliente['EditarCliente']).values('personaid')
    cliente_persona = Persona.objects.get(personaid=cliente[0]['personaid'])
    cliente_usuario = Usuario.objects.get(personaid=cliente[0]['personaid'])
    direccion1 = Direccioncliente.objects.filter(clienteid=editar_cliente['EditarCliente']).values('direccionid')
    direccion = Direccion.objects.get(direccionid=direccion1[0]['direccionid'])

    form1 = FormClienteNormal1(request.POST or None, instance=cliente_persona)
    form2 = FormClienteNormal2(request.POST or None, instance=direccion)
    form3 = FormClienteNormal3(request.POST or None, instance=cliente_usuario)

    if request.method == 'POST':

        run_cuerpo = request.POST.get('runcuerpo')
        dv = request.POST.get('dv')
        nombres = request.POST.get('nombres')
        apellido_paterno = request.POST.get('apellidopaterno')
        apellido_materno = request.POST.get('apellidomaterno')
        telefono = request.POST.get('telefono')
        email = request.POST.get('email')
        calle = request.POST.get('calle')
        numero = request.POST.get('numero')
        comuna_id = request.POST.get('comunaid')
        tipo_vivienda_id = request.POST.get('tipoviviendaid')
        tipo_barrio_id = request.POST.get('tipobarrioid')
        nombre_sector = request.POST.get('nombresector')
        # nombre_usuario = request.POST.get('nombreusuario')

        cliente_persona, created = Persona.objects.get_or_create(personaid=cliente[0]['personaid'])
        cliente_persona.runcuerpo = run_cuerpo
        cliente_persona.dv = dv
        cliente_persona.apellidopaterno = apellido_paterno
        cliente_persona.apellidomaterno = apellido_materno
        cliente_persona.nombres = nombres
        cliente_persona.telefono = telefono
        cliente_persona.save()

        cliente_direccion, created = Direccion.objects.get_or_create(direccionid=direccion1[0]['direccionid'])
        cliente_direccion.calle = calle
        cliente_direccion.numero = numero
        cliente_direccion.comunaid = Comuna.objects.get(comunaid=comuna_id)
        cliente_direccion.tipoviviendaid = Tipovivienda.objects.get(tipoviviendaid=tipo_vivienda_id)
        cliente_direccion.tipobarrioid = Tipobarrio.objects.get(tipobarrioid=tipo_barrio_id)
        cliente_direccion.nombresector = nombre_sector
        cliente_direccion.save()

        cliente_usuario, created = Usuario.objects.get_or_create(personaid=cliente[0]['personaid'])
        cliente_usuario.email = email
        cliente_usuario.save()

        user_django = User.objects.get(email=cliente_usuario)
        user_django.username = user_django.username
        user_django.first_name = nombres
        user_django.last_name = apellido_paterno
        user_django.email = email
        user_django.save()

        sweetify.success(request,"Cliente actualizado correctamente")
        return redirect('ver_cliente_vendedor')

    context = {
        'form1': form1,
        'form2': form2,
        'form3': form3,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'vendedores/editar_cliente_vendedor.html', context)

##********************Proveedores*******************************************************************

@login_required(login_url="ingreso")
def Agregar_proveedor(request):
    Seguimiento_paginas("Modulo Proveedores - Agregar Proveedor", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    form = FormProveedor()
    form2 = FormProveedorUsuario()
    form3 = FormDireProvee()

    if request.method == 'POST':
        rut_cuerpo = request.POST.get('rutcuerpo')
        dv = request.POST.get('dv')
        razonsocial = request.POST.get('razonsocial')
        telefono = request.POST.get('fono')
        rubro_id = request.POST.get('rubroid')
        email = request.POST.get('email')
        calle = request.POST.get('calle')
        numero = request.POST.get('numero')
        comuna_id = request.POST.get('comunaid')
        tipo_vivienda_id = request.POST.get('tipoviviendaid')
        tipo_barrio_id = request.POST.get('tipobarrioid')
        nombre_sector = request.POST.get('nombresector')
        nombre_usuario = request.POST.get('nombreusuario')
        contrase??a = request.POST.get('password')
        confirme_contrase??a = request.POST.get('confirme_contrase??a')

        proveedorRegistro = Proveedor.objects.filter(rutcuerpo=rut_cuerpo, dv=dv).exists()
        usuarioRegistro = Usuario.objects.filter(email=email).exists()
        usuarioRegistro3 = Usuario.objects.filter(nombreusuario=nombre_usuario).exists()

        if proveedorRegistro == True:
            sweetify.warning(request, "El rut ingresado ya existe")
        elif usuarioRegistro == True:
            sweetify.warning(request, "El correo ingresado ya existe")
        elif usuarioRegistro3 == True:
            sweetify.warning(request, "El usuario ingresado ya esta registrado")
        else:
            
            rolusuario = 5 #Rolusuario.objects.filter(descripcion="Proveedor").values('rolid')
            django_cursor = connection.cursor()
            cursor = django_cursor.connection.cursor()
            procedimiento = cursor.callproc(
                'SP_InsertProveedor', 
                [
                    int(rut_cuerpo), 
                    dv, 
                    razonsocial,
                    int(telefono), 
                    rubro_id,
                    calle,
                    str(numero),
                    nombre_sector,
                    tipo_vivienda_id,
                    tipo_barrio_id,
                    comuna_id,
                    email,
                    nombre_usuario,
                    contrase??a,
                    rolusuario
                ]
            )

            user = User.objects.create_user(
                username=nombre_usuario,
                first_name=razonsocial,
                last_name=razonsocial,
                email=email,
                is_superuser=False,
                is_active=True
            )
            user.set_password(contrase??a)
            user.set_password(confirme_contrase??a)

            if (procedimiento is not None and 
                user is not None):
                user.save()

                sweetify.success(request, "Se ha registrado correctamente")
                return redirect('index')
            else:
                sweetify.error(request, "No es posible registrarse en este momento")

    context = {
        'form': form,
        'form2': form2,
        'form3': form3,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'proveedores/agregar_proveedor.html', context)

def Cambiar_estado_proveedor(id_proveedor):
    proveedor = Proveedor.objects.get(proveedorid=id_proveedor)

    if proveedor.estadoid.descripcion == 'Activo':
        proveedor.estadoid = Estado.objects.get(descripcion="Inactivo")
        proveedor.save()
    else:
        proveedor.estadoid = Estado.objects.get(descripcion="Activo")
        proveedor.save()

@login_required(login_url="ingreso")
def Listar_orden(request):
    Seguimiento_paginas("Modulo Ordenes de Compra - Listar Ordenes", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo.rolid.descripcion
        proveid = tipo.proveedorid.proveedorid
    else: 
        tipo_usuario = None
        proveid = None

    if proveid is not None:
        ordenes = Ordencompra.objects.filter(proveedorid = proveid)
    else:
        context = {}
        return render(request, 'index.html', context)

    if request.method == 'POST':

        if request.POST.get('VerOrden') is not None:
            request.session['_ver_orden'] = request.POST
            return HttpResponseRedirect('ver_orden')

    context = {
        'ordenes': ordenes,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'proveedores/listar_ordenes.html', context)

@login_required(login_url="ingreso")
def Ver_orden(request):
    Seguimiento_paginas("Modulo Ordenes de Compra - Ver Detalle Orden", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    ver_orden = request.session.get('_ver_orden')
    orden = Ordencompra.objects.get(ordenid = ver_orden['VerOrden'])
    productos_orden = Detalleorden.objects.filter(ordenid = orden.ordenid)


    context = {
        'orden': orden,
        'tipo_usuario': tipo_usuario,
        'productos_orden': productos_orden
    }

    return render(request, 'proveedores/ver_orden.html', context)

@login_required(login_url="ingreso")
def Listar_proveedores(request):
    Seguimiento_paginas("Modulo Proveedores - Listar Proveedor", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    proveedor = Proveedor.objects.all()

    if request.method == 'POST':
        if request.POST.get('CambiarEstado') is not None:
            id_proveedor = request.POST.get('CambiarEstado')
            Cambiar_estado_proveedor(id_proveedor)
            proveedor = Proveedor.objects.get(proveedorid=id_proveedor)
            messages.warning(request,
                             f'El producto {proveedor.razonsocial} ha quedado {proveedor.estadoid.descripcion} correctamente')
            return redirect('listar_proveedores')

        if request.POST.get('VerProveedor') is not None:
            request.session['_old_post'] = request.POST
            return HttpResponseRedirect('ver_proveedor')

        if request.POST.get('EditarProveedor') is not None:
            request.session['_old_post'] = request.POST
            return HttpResponseRedirect('editar_proveedor')

    context = {
        'proveedor': proveedor,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'proveedores/listar_proveedores.html', context)


@login_required(login_url="ingreso")
def Ver_proveedor(request):
    Seguimiento_paginas("Modulo Proveedores - Ver Detalle Proveedor", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    old_post = request.session.get('_old_post')
    proveedor = Proveedor.objects.get(proveedorid=old_post['VerProveedor'])

    context = {
        'proveedor':proveedor,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'proveedores/ver_proveedor.html', context)


@login_required(login_url="ingreso")
def Editar_proveedor(request):
    Seguimiento_paginas("Modulo Proveedores - Editar Proveedor", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    old_post = request.session.get('_old_post')

    proveedor = Proveedor.objects.get(proveedorid=old_post['EditarProveedor'])
    form = FormProveedor(request.POST or None, instance=proveedor)

    if request.method == 'POST':
        razonsocial = request.POST.get('razonsocial')
        rutcuerpo = request.POST.get('rutcuerpo')
        dv = request.POST.get('dv')
        fono = request.POST.get('fono')
        rubroid = request.POST.get('rubroid')
        direccionid = request.POST.get('direccionid')
        estadoid = request.POST.get('estadoid')

        rubro = Tiporubro.objects.get(rubroid=rubroid)
        direccion = Direccion.objects.get(direccionid=direccionid)
        estado = Estado.objects.get(estadoid=estadoid)

        proveedor, created = Proveedor.objects.get_or_create(proveedorid=old_post['EditarProveedor'])
        try:
            proveedor.razonsocial = razonsocial
            proveedor.rutcuerpo = rutcuerpo
            proveedor.dv = dv
            proveedor.fono = fono
            proveedor.rubroid = rubro
            proveedor.direccionid = direccion
            proveedor.estadoid = estado
            proveedor.save()

            messages.warning(request, 'Proveedor actualizado correctamente')
            return redirect('listar_proveedores')

        except Exception as error:
            messages.error(request, error)

    context = {
        'form': form,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'proveedores/editar_proveedor.html', context)

# ********************************Empleados************************************************

@login_required(login_url="ingreso")
def Listar_empleados(request):
    Seguimiento_paginas("Modulo Empleados - Listar Empleados", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    empleados = Cargo.objects.get(descripcion="Empleado")
    empleados = Empleado.objects.filter(cargoid=empleados)

    if request.method == 'POST':

        if request.POST.get('CambiarEstado') is not None:
            id_empleado = request.POST.get('CambiarEstado')
            Cambiar_estado_empleado(id_empleado)
            empleado = Empleado.objects.get(empleadoid=id_empleado)
            sweetify.success(request,
                             f'El empleado {empleado.personaid.runcuerpo} - {empleado.personaid.dv} ha quedado {empleado.estadoid.descripcion} correctamente')
            return redirect('listar_empleados')

        if request.POST.get('VerEmpleado') is not None:
            request.session['_old_post'] = request.POST
            return HttpResponseRedirect('ver_empleado')

        if request.POST.get('EditarEmpleado') is not None:
            request.session['_old_post'] = request.POST
            return HttpResponseRedirect('editar_empleado')

    context = {
        'empleados':empleados,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'empleados/listar_empleados.html', context)

@login_required(login_url="ingreso")
def Agregar_empleado(request):
    Seguimiento_paginas("Modulo Empleados - Agregar Empleado", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    form1 = FormEmpleadoPersona()
    form2 = FormEmpleadoUsuario()
    form3 = FormEmpleadoEmpleado()

    if request.method == 'POST':

        run_cuerpo = request.POST.get('runcuerpo')
        dv = request.POST.get('dv')
        nombres = request.POST.get('nombres')
        apellido_paterno = request.POST.get('apellidopaterno')
        apellido_materno = request.POST.get('apellidomaterno')
        telefono = request.POST.get('telefono')
        nombre_usuario = request.POST.get('nombreusuario')
        fecha_ingreso = request.POST.get('fechaingreso')
        email = request.POST.get('email')
        contrase??a = request.POST.get('password')
        confirme_contrase??a = request.POST.get('confirme_contrase??a')

        fecha_ingreso = fecha_ingreso[6:10] + '-' + fecha_ingreso[3:5] + '-' + fecha_ingreso[0:2]

        personaRegistro = Persona.objects.filter(runcuerpo=run_cuerpo, dv=dv).exists()
        usuarioRegistro = Usuario.objects.filter(email=email).exists()
        usuarioRegistro3 = Usuario.objects.filter(nombreusuario=nombre_usuario).exists()

        if personaRegistro:
            sweetify.error(request, "El run ingresado ya existe")

        elif usuarioRegistro:
            sweetify.error(request, "El correo ingresado ya existe")

        elif usuarioRegistro3:
            sweetify.error(request, "El empleado ya cuenta con un usuario")

        else:
            cargo = Cargo.objects.filter(descripcion="Empleado").values('cargoid')
            rolusuario = Rolusuario.objects.filter(descripcion="Empleado").values('rolid')

            django_cursor = connection.cursor()
            cursor = django_cursor.connection.cursor()
            procedimiento = cursor.callproc(
                'SP_Insertempleado', 
                [
                    int(run_cuerpo), 
                    dv, 
                    apellido_paterno, 
                    apellido_materno, 
                    nombres, 
                    int(telefono), 
                    fecha_ingreso, 
                    int(cargo[0]['cargoid']), 
                    email, 
                    nombre_usuario, 
                    contrase??a, 
                    rolusuario[0]['rolid']
                ]
            )

            user = User.objects.create_user(
                username=nombre_usuario,
                first_name=nombres,
                last_name=apellido_paterno,
                email=email,
                is_superuser=False,
                is_active=True
            )
            user.set_password(contrase??a)
            user.set_password(confirme_contrase??a)

            if (procedimiento is not None and 
                user is not None):
                user.save()
                sweetify.success(request, "Empleado creado correctamente")
                return redirect('listar_empleados')
            else:
                sweetify.warning(request, "No es posible crear el empleado en este momento")

    context = {
        'form1': form1,
        'form2': form2,
        'form3': form3,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'empleados/agregar_empleado.html', context)

@login_required(login_url="ingreso")
def Ver_empleado(request):
    Seguimiento_paginas("Modulo Empleados - Ver Detalle Empleado", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    old_post = request.session.get('_old_post')
    empleado = Empleado.objects.get(empleadoid=old_post['VerEmpleado'])

    context = {
        'empleado':empleado,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'empleados/ver_empleado.html', context)

@login_required(login_url="ingreso")
def Editar_empleado(request):
    Seguimiento_paginas("Modulo Empleados - Editar Empleado", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    old_post = request.session.get('_old_post')

    empleado = Empleado.objects.filter(empleadoid=old_post['EditarEmpleado']).values('personaid')
    empleado_usuario = Usuario.objects.get(personaid=empleado[0]['personaid'])
    empleado_persona = Persona.objects.get(personaid=empleado[0]['personaid'])
    empleado_empleado = Empleado.objects.get(personaid=empleado[0]['personaid'])

    form1 = FormEmpleadoPersona(request.POST or None, instance=empleado_persona)
    form2 = FormEmpleadoUsuario(request.POST or None, instance=empleado_usuario)
    form3 = FormEmpleadoEmpleado(request.POST or None, instance=empleado_empleado)

    if request.method == 'POST':
        run_cuerpo = request.POST.get('runcuerpo')
        dv = request.POST.get('dv')
        nombres = request.POST.get('nombres')
        apellido_paterno = request.POST.get('apellidopaterno')
        apellido_materno = request.POST.get('apellidomaterno')
        telefono = request.POST.get('telefono')
        nombre_usuario = request.POST.get('nombreusuario')
        fecha_ingreso = request.POST.get('fechaingreso')
        email = request.POST.get('email')

        fecha_ingreso = fecha_ingreso[6:10] + '-' + fecha_ingreso[3:5] + '-' + fecha_ingreso[0:2]

        empleadoEm, created = Empleado.objects.get_or_create(empleadoid=old_post['EditarEmpleado'])
        empleadoEm.fechaingreso = fecha_ingreso
        empleadoEm.personaid.runcuerpo = run_cuerpo
        empleadoEm.personaid.dv = dv
        empleadoEm.personaid.nombres = nombres
        empleadoEm.personaid.apellidopaterno = apellido_paterno
        empleadoEm.personaid.apellidomaterno = apellido_materno
        empleadoEm.personaid.telefono = telefono
        empleadoEm.save()

        empleado_para_persona = Empleado.objects.filter(empleadoid=old_post['EditarEmpleado']).values('personaid')
        usuario, created = Usuario.objects.get_or_create(personaid=empleado_para_persona[0]['personaid'])
        usuario.nombreusuario = nombre_usuario
        usuario.email = email
        usuario.save()

        persona, created = Persona.objects.get_or_create(personaid=empleado_para_persona[0]['personaid'])
        persona.runcuerpo = run_cuerpo
        persona.dv = dv
        persona.apellidopaterno = apellido_paterno
        persona.apellidomaterno = apellido_materno
        persona.nombres = nombres
        persona.telefono = telefono
        persona.save()

        user_django = User.objects.get(email=empleado_usuario)
        user_django.username = user_django.username
        user_django.first_name = nombres
        user_django.last_name = apellido_paterno
        user_django.email = email
        user_django.save()

        sweetify.success(request, "Empleado actualizado correctamente")
        return redirect('listar_empleados')

    context = {
        'form1': form1,
        'form2': form2,
        'form3': form3,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'empleados/editar_empleado.html', context)


def Cambiar_estado_empleado(id_empleado):
    empleado = Empleado.objects.get(empleadoid=id_empleado)
    persona_id = Empleado.objects.filter(empleadoid=id_empleado).values('personaid')
    usuario = Usuario.objects.filter(personaid=persona_id[0]['personaid']).values('nombreusuario')
    user = User.objects.get(username=usuario[0]['nombreusuario'])

    if empleado.estadoid.descripcion == 'Activo':
        empleado.estadoid = Estado.objects.get(descripcion="Inactivo")
        empleado.save()
        user.is_active = False
        user.save()
    else:
        empleado.estadoid = Estado.objects.get(descripcion="Activo")
        empleado.save()
        user.is_active = True
        user.save()


# *********************************Pedidos************************************************
@csrf_exempt
def Crear_pedido(request, id=None):
    # Seguimiento_paginas("Modulo Ordenes de Compra - Crear Pedido", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    proveedor = Proveedor.objects.get(proveedorid=id)
    listaProds = Productoproveedor.objects.filter(proveedorid=proveedor)
    if request.method == 'POST':
        lista_productos = []
        for key, value in request.POST.items():

            try:
                Producto.objects.get(productoid=key)
                if len(lista_productos) == 0:
                    lista_productos.append({'id': key, 'value': value})
                else:
                    for product in lista_productos:
                        if product['id'] == key:
                            sum_value = int(product['value']) + int(value)
                            product.update({'value': sum_value})
                        else:
                            lista_productos.append({'id': key, 'value': value})
                            break

            except Exception as error:
                if len(value) == 0:
                    fecha = "1000-10-10"
                else:
                    partes = value.split("/")
                    fecha = "-".join(reversed(partes))
        proveedor_orden = Proveedor.objects.get(proveedorid=int(id))
        estado_orden = Estadoorden.objects.get(estadoordenid=1)

        ordenPedido = Ordencompra.objects.create(
            fechapedido=fecha,
            proveedorid=proveedor_orden,
            estadoordenid=estado_orden
        )

        last_orden_compra = Ordencompra.objects.order_by('ordenid').last()
        estado_detalle = Estado.objects.get(estadoid=1)

        for products in lista_productos:
            producto = Producto.objects.get(productoid=int(products['id']))

            detallePedido = Detalleorden.objects.create(
                productoid=producto,
                cantidad=products['value'],
                ordenid=last_orden_compra,
                estadoid=estado_detalle

            )
        sweetify.success(request, "Orden de pedido realizada con exito")
        return redirect('listar_pedidos')

    context = {
        'listaProds': listaProds,
        'tipo_usuario': tipo_usuario,
    }
    return render(request, 'pedidos/crear_pedido.html', context)

@login_required(login_url="ingreso")
def filtro_proveedor(request):

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    proveedores = Proveedor.objects.all()
    context = {
        'proveedores': proveedores,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'pedidos/crear_pedido_proveedores.html', context)

@login_required(login_url="ingreso")
def cambiar_estado_pedido(id_pedido):
    orden_compra = Ordencompra.objects.get(ordenid=int(id_pedido))
    detalle_orden = Detalleorden.objects.filter(ordenid=int(id_pedido))
    estado_eliminado = Estadoorden.objects.get(estadoordenid=23)
    estado_eliminado_detale = Estado.objects.get(estadoid=2)

    if orden_compra.estadoordenid.estadoordenid == 1:
        orden_compra.estadoordenid = estado_eliminado
        orden_compra.save()

    for detalle in detalle_orden:
        detalle.estadoid = estado_eliminado_detale
        detalle.save()

@login_required(login_url="ingreso")
def Listar_pedidos(request):
    Seguimiento_paginas("Modulo Ordenes de Compra - Listar Pedidos", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    ordenes = Ordencompra.objects.all()

    if request.method == 'POST':
        if request.POST.get('CambiarEstado') is not None:
            id_pedido = request.POST.get('CambiarEstado')
            cambiar_estado_pedido(id_pedido)
            orden_compra = Ordencompra.objects.get(ordenid=id_pedido)
            sweetify.success(request,
                             f'La Nro:{orden_compra.ordenid} se elimino correctamente')
            return redirect('listar_pedidos')

        if request.POST.get('VerPedido') is not None:
            request.session['_old_post'] = request.POST
            return HttpResponseRedirect('ver_pedido')

    context = {
        'ordenes': ordenes,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'pedidos/listar_pedidos.html', context)

@login_required(login_url="ingreso")
def Ver_pedidos(request):
    Seguimiento_paginas("Modulo Ordenes de Compra - Ver Detalle Pedido", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    old_post = request.session.get('_old_post')
    detalle_orden = Detalleorden.objects.filter(ordenid=old_post['VerPedido'])
    recepcion_orden = Recepcion.objects.filter(ordenid=old_post['VerPedido'])

    context = {
        'detalle_orden': detalle_orden,
        'recepcion_orden':recepcion_orden,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'pedidos/ver_pedido.html', context)

# ***************************Recepcionar pedidos*******************************************

@login_required(login_url="ingreso")
def RecepcionPedido(request, id=None):
    Seguimiento_paginas("Modulo Recepcion - Recepcionar un pedido", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    if id:
        orden_pedido = Ordencompra.objects.filter(ordenid=id)
        detalle_orden = Detalleorden.objects.filter(ordenid=id, estadoid=1)
        productos = Producto.objects.all()

    if request.method == 'POST':
        for key, value in request.POST.items():
            orden_pedido = Ordencompra.objects.get(ordenid=id)
            detalle_orden = Detalleorden.objects.filter(ordenid=orden_pedido)
            if key == "csrfmiddlewaretoken":
                pass
            else:
                try:
                    orden_pedido = Ordencompra.objects.get(ordenid=id)
                    producto_r = Producto.objects.get(productoid=key)
                    id_producto = producto_r.productoid
                    detalle = Detalleorden.objects.get(ordenid=id, productoid=id_producto)

                    Recepcion.objects.create(
                        fecharecepcion=datetime.now().date(),
                        cantidad=detalle.cantidad,
                        productoid=detalle.productoid,
                        proveedorid=orden_pedido.proveedorid,
                        ordenid=orden_pedido
                    )

                    detalle, created = Detalleorden.objects.get_or_create(ordenid=id, productoid=id_producto)
                    detalle.estadoid = Estado.objects.get(estadoid=2)
                    detalle.save()

                    producto_recepcionar, created = Producto.objects.get_or_create(productoid=id_producto)
                    producto_recepcionar.stock = producto_recepcionar.stock + detalle.cantidad
                    producto_recepcionar.save()
                    sweetify.success(request, "Productos recepcionados satisfactoriamente")
                except Exception as error:
                    sweetify.warning(request, "No es posible Recepcionar el producto en este momento")

        verificar_detalle_orden = Detalleorden.objects.filter(ordenid=id, estadoid=1)
        if len(verificar_detalle_orden) == 0:
            orden, created = Ordencompra.objects.get_or_create(ordenid=id)
            orden.estadoordenid = Estadoorden.objects.get(estadoordenid=22)
            orden.save()
        else:
            orden, created = Ordencompra.objects.get_or_create(ordenid=id)
            orden.estadoordenid = Estadoorden.objects.get(estadoordenid=41)
            orden.save()

        return redirect('listar_pedidos')

    context = {
        'ordenPedido': orden_pedido,
        'productos': productos,
        'detalleOrden': detalle_orden,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'recepcion_pedido.html', context)

#****************************Boletas*******************************************************
@login_required(login_url="ingreso")
def Listar_boletas(request):
    Seguimiento_paginas("Modulo Boletas - Listar Boletas", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    boletas = Boleta.objects.all()

    if request.method == 'POST':

        if request.POST.get('VerBoleta') is not None:
            request.session['_ver_boleta'] = request.POST
            return HttpResponseRedirect('ver_boleta')

    context = {
        'boletas': boletas,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'boletas/listar_boletas.html', context)

@login_required(login_url="ingreso")
def Ver_boleta(request):
    Seguimiento_paginas("Modulo Boletas - Ver Detalle Boleta", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    ver_boleta = request.session.get('_ver_boleta')

    boleta = Boleta.objects.get(nroboleta = ver_boleta['VerBoleta'])
    productos_boleta = Detalleventa.objects.filter(nroventa = boleta.nroventa)
    venta = Venta.objects.get(nroventa = boleta.nroventa.nroventa)
    direccion_cliente = Direccioncliente.objects.get(clienteid=venta.clienteid)
    if venta.clienteid.personaid != None:
        giro = "persona Natural"
    else:
        giro = "Distribuidor Ferreteria"

    if request.method == 'POST':

        tipoInforme = request.POST.get('informeCheck')
        descargarInforme = request.POST.get('descargarInforme')
        
        if tipoInforme is not None and descargarInforme is not None:

            lista1 = []
            lista2 = []
            lista3 = []
            lista1.append(["Nro Boleta","Fecha","Total","Estado"])  
            val = Boleta.objects.filter(nroboleta = ver_boleta['VerBoleta']).values_list('nroboleta','fechaboleta','totalboleta','estadoid__descripcion')

            for valores in val:

                lista1.append(list(valores)) 

            lista2.append(["Nro Venta","Tipo pago","Run cliente","DV"])  
            val = Boleta.objects.filter(nroboleta = ver_boleta['VerBoleta']).values_list('nroventa__nroventa','nroventa__tipodocumentoid__descripcion','nroventa__clienteid__personaid__runcuerpo','nroventa__clienteid__personaid__dv')

            for valores in val:

                lista2.append(list(valores))


            for valores in val:

                lista3.append(list(valores))

            if tipoInforme == "informeExcel":

                nombre_archivo = "Detalle boleta"
                tipo_doc = 'ms-excel'
                extension = 'xlsx'
                
                lista1.append("")
                lista2.append("")

                lista = lista1 + lista2 + lista3
                
                return  creacion_excel(nombre_archivo, lista, tipo_doc, extension)

            if tipoInforme == "informePdf":
                
                response = HttpResponse(content_type=f'application/pdf')  
                buff = BytesIO()  

                c = canvas.Canvas(buff, pagesize=letter)
                c= generar_factura(c, venta, boleta, productos_boleta, direccion_cliente, giro, 0)
                c.showPage()
                c.save()

                response.write(buff.getvalue())   
                buff.seek(0)

                return FileResponse(buff, as_attachment=False, filename=f'factura.pdf')

            if tipoInforme == "informeWord":

                nombre_archivo = "Detalle boleta"
                # return  creacion_doc(lista1, nombre_archivo)
                document = Document()
                document.add_heading('Detalle Boleta', 0)

                filas = 0
                for x in lista1:
                    columnas = len(x)
                    filas += 1

                # add grid table
                table = document.add_table(rows=filas, cols=columnas, style="Table Grid")

                for x in range(columnas):
                    table.rows[0].cells[x]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w'))))
                
                # access first row's cells
                heading_row = table.rows[0].cells

                # add headings
                cont = 0
                for value in lista1[0]:
                    heading_row[cont].text = value
                    cont += 1

                lista1.pop(0)
                cont = 0
                cont2 = 0

                for value in lista1:
                    cont += 1
                    data_row = table.rows[cont].cells

                    for x in value:
                        data_row[cont2].text = f'{x}'
                        cont2 += 1
                    cont2 = 0

                document.add_paragraph("")

                # parrafo.add_run().add_break()
                filas = 0
                for x in lista2:
                    columnas = len(x)
                    filas += 1

                # add grid table
                table2 = document.add_table(rows=filas, cols=columnas, style="Table Grid")

                for x in range(columnas):
                    table2.rows[0].cells[x]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w'))))

                # access first row's cells
                heading_row = table2.rows[0].cells

                # add headings
                cont = 0
                for value in lista2[0]:
                    heading_row[cont].text = value
                    cont += 1

                lista2.pop(0)
                cont = 0
                cont2 = 0

                for value in lista2:
                    cont += 1
                    data_row = table2.rows[cont].cells

                    for x in value:
                        data_row[cont2].text = f'{x}'
                        cont2 += 1
                    cont2 = 0 

                document.add_paragraph("")

                filas = 0
                for x in lista3:
                    columnas = len(x)
                    filas += 1

                # add grid table
                table3 = document.add_table(rows=filas, cols=columnas, style="Table Grid")

                for x in range(columnas):
                    table3.rows[0].cells[x]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w'))))

                # access first row's cells
                heading_row = table3.rows[0].cells

                # add headings
                cont = 0
                for value in lista3[0]:
                    heading_row[cont].text = value
                    cont += 1

                lista3.pop(0)
                cont = 0
                cont2 = 0

                for value in lista3:
                    cont += 1
                    data_row = table3.rows[cont].cells

                    for x in value:
                        data_row[cont2].text = f'{x}'
                        cont2 += 1
                    cont2 = 0  

                response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename={nombre_archivo}.docx'
                document.save(response)

                return response

    context = {
        'boleta': boleta,
        'tipo_usuario': tipo_usuario,
        'productos_boleta': productos_boleta
    }

    return render(request, 'boletas/ver_boleta.html', context)

def Cambiar_estado_boleta(id_boleta):

    boleta = Boleta.objects.get(nroboleta=id_boleta)

    if boleta.estadoid.descripcion == 'Activo':
        boleta.estadoid = Estado.objects.get(descripcion="Inactivo")
        boleta.save()
    else:
        boleta.estadoid = Estado.objects.get(descripcion="Activo")
        boleta.save()


#****************************Despacho*******************************************************
@login_required(login_url="ingreso")
def Listar_despacho(request):
    Seguimiento_paginas("Modulo Despachos - Listar Despachos", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    despachos = Despacho.objects.all().order_by('despachoid')


    if request.method == 'POST':

        if request.POST.get('VerDespacho') is not None:
            request.session['_ver_despacho'] = request.POST
            return HttpResponseRedirect('ver_despacho')
        

    if request.method == 'POST':
        estado = request.POST.get('estados')

    context = {
        'despachos': despachos,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'despacho/listar_despacho.html', context)

@login_required(login_url="ingreso")
def Ver_despacho(request):
    Seguimiento_paginas("Modulo Despachos - Ver Detalle Despacho", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    ver_despacho = request.session.get('_ver_despacho')
    despacho = Despacho.objects.get(despachoid = ver_despacho['VerDespacho'])
    productos_despacho = Detalleventa.objects.filter(nroventa = despacho.nroventa.nroventa)
    venta = Venta.objects.get(nroventa = despacho.nroventa.nroventa)
    direccion_cliente = Direccioncliente.objects.get(clienteid= venta.clienteid)
    
    if despacho.tipodespacho == 'Despacho':
        guia = Guiadespacho.objects.get(despachoid = despacho.despachoid)

    if venta.clienteid.personaid != None:
        giro = "persona natural"
    else:
        giro = "Distribuidor Ferreteria"


    if request.method == 'POST':

        if request.POST.get('GeneraGuia'):
            response = HttpResponse(content_type=f'application/pdf')  
            buff = BytesIO()  
            c = canvas.Canvas(buff, pagesize=letter)
            c= generar_factura(c, venta, guia, productos_despacho, direccion_cliente, giro, 2)
            c.showPage()
            c.save()

            response.write(buff.getvalue())   
            buff.seek(0)

            return FileResponse(buff, as_attachment=False, filename=f'boleta.pdf')
            
        else:
            response.write(buff.getvalue())   
            buff.seek(0)
            despachoId = request.POST.get('despacho_id')
            estado_id = request.POST.get('estado_id')
            estado = request.POST.get('btnAccion')
            Cambiar_estado_despacho(despachoId,estado_id,estado)
            return redirect('listar_despacho')        

    context = {
        'despacho': despacho,
        'tipo_usuario': tipo_usuario,
        'productos_despacho': productos_despacho
    }

    return render(request, 'despacho/ver_despacho.html', context)

def Cambiar_estado_despacho(despachoId,estado_id,estado):
    despacho = Despacho.objects.get(despachoid=despachoId)
    if estado == 'Cancelar':
        estadoActual = 'Inactivo'
    else:
        estadoActual = 'Despachado'

    if despacho.estadoid.descripcion == 'Activo':
        despacho.estadoid = Estado.objects.get(descripcion=estadoActual)
        despacho.save()

    else:
        despacho.estadoid = Estado.objects.get(descripcion="Activo")
        despacho.save()

#****************************Orden Proveedor******************************************
@login_required(login_url="ingreso")
def Ver_Orden(request):
    Seguimiento_paginas("Modulo Despachos - Ver Despacho", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    ver_despacho = request.session.get('_ver_despacho')
    despacho = Despacho.objects.get(despachoid = ver_despacho['VerDespacho'])
    productos_despacho = Detalleventa.objects.filter(nroventa = despacho.nroventa.nroventa)

    if request.method == 'POST':

        despachoId = request.POST.get('despacho_id')
        estado_id = request.POST.get('estado_id')
        estado = request.POST.get('btnAccion')
        Cambiar_estado_despacho(despachoId,estado_id,estado)
        return redirect('listar_despacho')

    context = {
        'despacho': despacho,
        'tipo_usuario': tipo_usuario,
        'productos_despacho': productos_despacho
    }

    return render(request, 'proveedor/ver_orden.html', context)

#****************************Creacion de archivos******************************************

def creacion_excel(nombre_archivo, lista, tipo_doc, extension):
    book = openpyxl.Workbook()  # Se crea un libro
    sheet = book.active  # Se activa la primera hojar
    sheet.title = f"{nombre_archivo}"
    cont = 0
    cont2 = 0

    for i in lista:
        cont += 1
        for j in i:
            
            cont2 += 1
            val = sheet.cell(row=cont, column=cont2)
            val.value = j
        cont2 = 0    

    cont = 0
    cont2 = 0
    
    response = HttpResponse(content_type=f"application/{tipo_doc}")
    contenido = "attachment; filename = {0}.{1}".format(nombre_archivo, extension)
    response["Content-Disposition"] = contenido
    book.save(response)

    return response

def creacion_pdf(lista,tipo_doc,tama??o_pagina, nombre, extension,lista2 = None, valor = None):
    response = HttpResponse(content_type=f'application/{tipo_doc}')  

    buff = BytesIO()  

    doc = SimpleDocTemplate(buff,  
        pagesize=tama??o_pagina,  
        rightMargin=40,  
        leftMargin=40,  
        topMargin=60,  
        bottomMargin=18,  
    ) 
    
    data = []  
    styles = getSampleStyleSheet()  
    styles = styles['Heading1']
    styles.alignment = TA_CENTER 

    header = Paragraph(f"{nombre}", styles)  
    
    data.append(header)  

    t = Table(lista)  

    t.setStyle(TableStyle(  
        [  
        ('GRID', (0, 0), (12, -1), 1, colors.dodgerblue),  
        ('LINEBELOW', (0, 0), (-1, 0), 3, colors.darkblue),  
        ('BACKGROUND', (0, 0), (-1, 0), colors.dodgerblue)  
        ]   
    ))  
    
    data.append(t)

    if lista2:

        styles = getSampleStyleSheet()  
        styles.pageBreakBefore = 2
        styles = styles['Heading1']
        styles.alignment = TA_CENTER 
        header = Paragraph(f" ", styles)  
        
        data.append(header)  

        t = Table(lista2)  

        t.setStyle(TableStyle(  
            [  
            ('GRID', (0, 0), (12, -1), 1, colors.dodgerblue),  
            ('LINEBELOW', (0, 0), (-1, 0), 3, colors.darkblue),  
            ('BACKGROUND', (0, 0), (-1, 0), colors.dodgerblue)  
            ]  
        ))  
        
        data.append(t)

    doc.build(data)  
    response.write(buff.getvalue())   

    buff.seek(0)

    respuesta = FileResponse(buff, as_attachment=valor, filename=f'{nombre}.{extension}')

    return respuesta

def creacion_doc(lista, nombre_archivo):
    document = Document()
    document.add_heading('Document Title', 0)

    filas = 0
    for x in lista:
        columnas = len(x)
        filas += 1

    # add grid table
    table = document.add_table(rows=filas, cols=columnas, style="Table Grid")

    # access first row's cells
    heading_row = table.rows[0].cells

    # add headings
    cont = 0
    for value in lista[0]:
        heading_row[cont].text = value
        cont += 1

    lista.pop(0)
    cont = 0
    cont2 = 0

    for value in lista:
        cont += 1
        data_row = table.rows[cont].cells

        for x in value:
            data_row[cont2].text = f'{x}'
            cont2 += 1
        cont2 = 0     

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename={nombre_archivo}.docx'
    document.save(response)

    return response

registerFont(TTFont('Arial','arial.ttf'))
# tipo_tributario=0 BOLETA
# tipo_tributario=1 FACTURA
def generar_factura(c, venta, documento, detalle_venta, direccion_cliente, giro, tipo_tributario='1'):
    c.translate(inch,inch)
    # define a large font
    c.setFont("Helvetica", 14)
    # choose some colors
    c.setStrokeColorRGB(0.1,0.8,0.1)
    c.setFillColorRGB(0,0,1) # font colour

    # c.drawImage('src\static\images\Ferme-logo.jpg',0*inch,8.7*inch, width=70, height=70)

    c.setFillColorRGB(255,0,0) # font colour
    c.setFont("Arial", 18)
    c.drawString(1.2*inch, 9.3*inch, "FERME SPA")
    c.setFont("Arial", 13)
    c.setFillColorRGB(0,0,255) # font colour
    c.drawString(1.2*inch, 9.1*inch, "Giro: Ferreteria ")
    c.drawString(1.2*inch, 8.9*inch, "San Bernardo - Santiago")
    #linea verde
    c.setStrokeColorCMYK(0,0,0,1) # vertical line colour 
    
    dt = date.today().strftime('%d-%b-%Y')

    #cuadrado arriba - derecha
    c.setStrokeColorCMYK(0,0,0,1) # Color linea 
    c.line(4*inch,9.8*inch,4*inch,8.6*inch)# linea vertical 
    c.line(6.8*inch,9.8*inch,6.8*inch,8.6*inch)# linea vertical 
    c.line(6.8*inch,9.8*inch,4*inch,9.8*inch)# linea horizontal  
    c.line(6.8*inch,8.6*inch,4*inch,8.6*inch)# linea horizontal  

    #info dentro de cuadrado
    c.setFillColorRGB(1,0,0) # font colour
    c.setFont("Arial", 14)
    c.drawString(4.6*inch,9.5*inch,'RUT: 69.736.194-2')
   
    c.setFont("Arial", 12)
    c.drawString(4.2*inch,8.1*inch,f'Fecha Emision:  {dt}')

    #BAJO EL LOGO
    c.setStrokeColorCMYK(0,0,0,1) # Color linea 
    c.line(0*inch,8.5*inch,0*inch,7.4*inch)# linea vertical 
    c.line(4*inch,8.5*inch,4*inch,7.4*inch)# linea vertical 
    c.line(0*inch,8.5*inch,4*inch,8.5*inch)# linea horizontal  

    c.setFont("Arial", 11)
    c.setFillColorRGB(0,0,255) # font colour
    c.drawString(0.1*inch, 8.3*inch, "SE??OR(ES):")
    c.drawString(0.1*inch, 8.1*inch, "GIRO:")
    c.drawString(0.1*inch, 7.9*inch, "DIRECCION:")
    c.drawString(0.1*inch, 7.7*inch, "COMUNA:")
    c.drawString(0.1*inch, 7.5*inch, "CONTACTO:")
    if tipo_tributario == 0: #BOLETA

        c.setFillColorRGB(1,0,0) # font colour
        c.setFont("Arial", 14)
        c.drawString(4.2*inch,9.1*inch,'BOLETA ELECTRONICA')
        c.drawString(5*inch,8.7*inch,f'NRO?? {documento.nroboleta}')

        c.setFillColorRGB(0,0,0) # font colour
        c.setFont("Arial", 11)

        c.drawString(1.1*inch, 7.5*inch, f"{str(venta.clienteid.personaid.telefono)}")
        c.drawString(1.05*inch, 8.3*inch, f"{str(venta.clienteid).lower().capitalize()}")
        
        neto = round(int(documento.totalboleta)*0.81)
        iva = round(int(documento.totalboleta)*0.19)

        c.drawRightString(6.5*inch,1.7*inch,f'{neto}') #  
        c.drawRightString(6.5*inch,1.5*inch,f'{iva}') # Total 
        c.drawRightString(6.5*inch,1.3*inch,f'{str(documento.totalboleta)}') # Total 

    elif tipo_tributario == 1: #FACTURA
        c.setFillColorRGB(1,0,0) # font colour
        c.setFont("Arial", 14)
        c.drawString(4.2*inch,9.1*inch,'FACTURA ELECTRONICA')
        c.drawString(5*inch,8.7*inch,f'NRO?? {documento.numerofactura}')

        c.setFillColorRGB(0,0,0) # font colour
        c.setFont("Arial", 11)
        c.drawString(1.1*inch, 7.5*inch, f"{str(venta.clienteid.personaid.telefono)}")
        c.drawString(1.05*inch, 8.3*inch, f"{str(venta.clienteid).lower().capitalize()}")

        c.drawRightString(6.5*inch,1.7*inch,f'{round(documento.neto)}') # Total 
        c.drawRightString(6.5*inch,1.5*inch,f'{round(documento.iva)}') # Total 
        c.drawRightString(6.5*inch,1.3*inch,f'{str(documento.totalfactura)}') # Total 

    elif tipo_tributario == 2: #GUIA
        c.setFillColorRGB(1,0,0) # font colour
        c.setFont("Arial", 14)
        c.drawString(4.2*inch,9.1*inch,'   GUIA DESPACHO')
        c.drawString(5*inch,8.7*inch,f'NRO?? {documento.nroguia}')

        c.setFillColorRGB(0,0,0) # font colour
        c.setFont("Arial", 11)
        c.drawString(1.1*inch, 7.5*inch, f"{str(venta.clienteid.personaid.telefono)}")
        c.drawString(1.05*inch, 8.3*inch, f"{str(venta.clienteid).lower().capitalize()}")

        try:
            documento_venta = Factura.objects.get(nroventa=venta.nroventa)
            c.drawRightString(6.5*inch,1.7*inch,f'{round(documento_venta.neto)}') # Total 
            c.drawRightString(6.5*inch,1.5*inch,f'{round(documento_venta.iva)}') # Total 
            c.drawRightString(6.5*inch,1.3*inch,f'{str(documento_venta.totalfactura)}') # Total 
        except Factura.DoesNotExist:
            documento_venta = Boleta.objects.get(nroventa=venta.nroventa)

            neto = str(int(documento_venta.totalboleta)-(int(documento_venta.totalboleta)*0.19))
            iva = str(int(documento_venta.totalboleta)*0.19)

            c.drawRightString(6.5*inch,1.7*inch,f'{neto}') #  
            c.drawRightString(6.5*inch,1.5*inch,f'{iva}') # Total 
            c.drawRightString(6.5*inch,1.3*inch,f'{str(documento_venta.totalboleta)}') # Total 

    elif tipo_tributario == 3: #NOTA CREDITO
        c.setFillColorRGB(1,0,0) # font colour
        c.setFont("Arial", 14)
        c.drawString(4.5*inch,9.1*inch,'NOTA DE CREDITO')
        c.drawString(5*inch,8.7*inch,f'NRO?? {documento.nronota}')

        c.setFillColorRGB(0,0,0) # font colour
        c.setFont("Arial", 11)
        c.drawString(1.1*inch, 7.5*inch, f"{str(venta.clienteid.personaid.telefono)}")
        c.drawString(1.05*inch, 8.3*inch, f"{str(venta.clienteid).lower().capitalize()}")

        neto = str(int(documento.total)-(int(documento.total)*0.19))
        iva = str(int(documento.total)*0.19)

        c.drawRightString(6.5*inch,1.7*inch,f'{neto}') #  
        c.drawRightString(6.5*inch,1.5*inch,f'{iva}') # Total 
        c.drawRightString(6.5*inch,1.3*inch,f'{str(documento.total)}') # Total 

    c.drawString(0.6*inch, 8.1*inch, f"{giro.lower().capitalize()}")
    c.drawString(1.1*inch, 7.9*inch, f"{str(direccion_cliente.direccionid).lower().capitalize()}")
    c.drawString(0.9*inch, 7.7*inch, f"{str(direccion_cliente.direccionid.comunaid).lower().capitalize()}")

    # c.rotate(-45) # restore the rotation 
    c.setFillColorRGB(0,0,0) # font colour
    c.setFont("Arial", 16)
    c.drawString(0.5*inch,7.2*inch,'Descripcion')
    c.drawString(4*inch,7.2*inch,'Precio')
    c.drawString(4.9*inch,7.2*inch,'Cantidad')
    c.drawString(6*inch,7.2*inch,'Total')

    # TABLA PRODUCTOS
    c.setStrokeColorCMYK(0,0,0,1) # vertical line colour 
    c.line(0,7.4*inch,6.8*inch,7.4*inch)
    c.line(0,7.1*inch,6.8*inch,7.1*inch)

    c.line(0*inch,7.4*inch,0*inch,2.5*inch)# first vertical line
    c.line(3.9*inch,7.4*inch,3.9*inch,2.5*inch)# second vertical line
    c.line(4.8*inch,7.4*inch,4.8*inch,2.5*inch)# third vertical line
    c.line(5.9*inch,7.4*inch,5.9*inch,2.5*inch)# fourty vertical line
    c.line(6.8*inch,7.4*inch,6.8*inch,2.5*inch)# fifty vertical line
    c.line(0.01*inch,2.5*inch,6.8*inch,2.5*inch)# horizontal line total

    #CAJA MONTOS    
    c.setStrokeColorCMYK(0,0,0,1) # Color linea 
    c.line(4*inch,1*inch,4*inch,2*inch)# linea vertical 
    c.line(6.8*inch,1*inch,6.8*inch,2*inch)# linea vertical 
    c.line(6.8*inch,1*inch,4*inch,1*inch)# linea horizontal  
    c.line(6.8*inch,2*inch,4*inch,2*inch)# linea horizontal  

    c.setFillColorRGB(0,0,0) # font colour
    c.setFont("Arial", 13)
    c.drawRightString(5.5*inch,1.7*inch,'Monto Neto $') # Total 
    c.drawRightString(5.5*inch,1.5*inch,'I.V.A 19% $') # Total 
    c.drawRightString(5.5*inch,1.3*inch,'TOTAL $') # Total 

    c.setStrokeColorRGB(0,0,0) # Bottom Line colour 

    c.line(0*inch,0.5*inch,0*inch,-0.6*inch)# first vertical line
    c.line(6.8*inch,0.5*inch,6.8*inch,-0.6*inch)# fifty vertical line
    c.line(0,-0.6*inch,6.8*inch,-0.6*inch)
    c.line(0, 0.5*inch,6.8*inch,0.5*inch)

    c.setFillColorRGB(0,0,0) # font colour
    c.setFont("Arial", 10)
    c.drawRightString(0.7*inch,0.2*inch,'NOMBRE:') # Total 
    c.drawRightString(2.5*inch,0.2*inch,'RUT:') # Total 
    c.drawRightString(4*inch,0.2*inch,'FECHA:') # Total 
    c.drawRightString(6*inch,0.2*inch,'RECINTO:') # Total 

    c.drawRightString(0.7*inch,0.0*inch,'FIRMA: ') # Total 

    c.setFont("Arial", 7.8)
    c.drawRightString(6.7*inch,-0.2*inch,'"El acuso de recibo que se desidira en este acto, de acuerdo a lo dispuesto en la letra b) de Art 4??, y la letra c) del Art 5?? de la ley 19.983,') # Total 
    c.drawRightString(4.3*inch,-0.35*inch,'acredita que la entrega de marcadores o servicio(s) prestado(s) ha(s) sido recibido(s)"') # Total 

    #PIE DE PAGINA
    c.setFont("Arial", 8) # font size
    c.setFillColorRGB(1,0,0) # font colour
    c.drawString(0, -0.9*inch,"www.ferme.cl")
    c.drawRightString(6.85*inch, -0.9*inch,"CEDIBLE")

    c.setFillColorRGB(0,0,1) # font colour
    c.setFont("Helvetica", 13)
    row_gap=0.2
    line_y=6.8 

    c.setFillColorRGB(0,0,0) # font colour
    c.setFont("Arial", 11) # font size

    for producto in detalle_venta:
        c.drawString(0.1*inch,line_y*inch,str(producto.productoid).lower().capitalize()) # p Name
        c.drawRightString(4.5*inch,line_y*inch,f'${str(producto.productoid.precio)}') # p Price
        c.drawRightString(5.5*inch,line_y*inch,str(producto.cantidad)) # p Qunt 
        c.drawRightString(6.5*inch,line_y*inch,f'${str(producto.subtotal * producto.cantidad)}') # Sub Total 
        line_y=line_y-row_gap

    return c


#************************************Ventas**************************************

@login_required(login_url="ingreso")
def crear_venta(request):
    Seguimiento_paginas("Modulo Ventas - Crear Venta", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    productos = Productoproveedor.objects.all().exclude(productoid__stock__lte=0)
    form = FormCliente()
    form_doc = FormVenta()
    form_docu =FormDocu()

    total = 0
    if request.method == 'POST':
        productos_venta = []
        for key, value in request.POST.items():
            if key == "valor_total" and value == "":
                sweetify.warning(request, "Porfavor ingrese productos")
            elif value == "":
                sweetify.warning(request, "Ingrese todo los datos")
            elif key == "csrfmiddlewaretoken":
                pass
            elif key == "clienteid":
                cliente_venta = Cliente.objects.get(clienteid=value)
            elif key == "tipodocumentoid":
                documento_venta = Tipodocumento.objects.get(tipodocumentoid=value)
            elif key == "tipopagoid":
                tipopagoid = Tipopago.objects.get(tipopagoid=value)
            elif key == "tipo_entrega":
                tipo_entrega = value

            else:
                try:
                    key = int(key)
                    try:
                        producto = Producto.objects.get(productoid=int(key))
                    except Producto.DoesNotExist:
                        sweetify.warning(request, "Ocurrio un Problema")
                        return render(request, 'index.html')
                except Exception as error:
                    if "cantidad" in key:
                        cantidad = value
                    elif key != "valor_total" and "total" in key:
                        total = value.replace("$", "")
                        productos_venta.append([producto, cantidad, total])
                    elif key == "valor_total":
                        total_venta = value
        if len(productos_venta) >0: 
            Venta.objects.create(
                fechaventa=datetime.now().date(),
                totalventa=total_venta,
                tipodocumentoid=documento_venta,
                clienteid=cliente_venta,
                tipopagoid=tipopagoid 
            )

            ultima_ventas = Venta.objects.order_by('nroventa').last()
            for producto in productos_venta:

                Detalleventa.objects.create(
                    cantidad = producto[1],
                    subtotal = int(producto[2]) * int(producto[1]),
                    productoid = producto[0],
                    nroventa = ultima_ventas
                )
                producto_modificar = producto[0]
                producto_vendido, created = Producto.objects.get_or_create(productoid=producto_modificar.productoid)
                producto_vendido.stock = producto_vendido.stock-int(producto[1])
                producto_vendido.save()

            if documento_venta.tipodocumentoid == 2:
                Factura.objects.create(
                    fechafactura = datetime.now().date(),
                    neto = int(total_venta)*0.81,
                    iva = int(total_venta)*0.19,
                    totalfactura = total_venta,
                    nroventa = ultima_ventas,
                    estadoid = Estado.objects.get(descripcion="Activo")
                )

                ultima_factura = Factura.objects.order_by('numerofactura').last()
                detalle_venta = Detalleventa.objects.filter(nroventa = ultima_factura.nroventa)
                direccion_cliente = Direccioncliente.objects.get(clienteid=ultima_ventas.clienteid)
                if ultima_ventas.clienteid.personaid != None:
                    giro = "persona natural"
                else:
                    giro = "Distribuidor Ferreteria"

                if int(tipo_entrega) == 1:
                    Despacho.objects.create(
                        fechasolicitud = datetime.now().date(),
                        fechadespacho =  datetime.now().date(),
                        nroventa = ultima_ventas,
                        estadoid = Estado.objects.get(descripcion="Activo"),
                        tipodespacho = "Despacho"
                    )

                    ultimo_despacho = Despacho.objects.order_by('despachoid').last()
                    direccion_cliente = Direccioncliente.objects.get(clienteid=cliente_venta)
                    Guiadespacho.objects.create(
                        fechaguia = datetime.now().date(),
                        despachoid = ultimo_despacho,
                        iddircliente = direccion_cliente
                    )
                else: 
                    Despacho.objects.create(
                        fechasolicitud = datetime.now().date(),
                        fechadespacho =  datetime.now().date(),
                        nroventa = ultima_ventas,
                        estadoid = Estado.objects.get(descripcion="Activo"),
                        tipodespacho = "Retiro")
                      
                messages.warning(request, 'Venta realizada con exito')
                productos_venta.insert(0 , ["Nombre Producto", "Cantidad", "Total"]) 

                response = HttpResponse(content_type=f'application/pdf')  
                buff = BytesIO()  

                c = canvas.Canvas(buff, pagesize=letter)
                c= generar_factura(c, ultima_ventas, ultima_factura, detalle_venta, direccion_cliente, giro, 1)
                c.showPage()
                c.save()

                response.write(buff.getvalue())   
                buff.seek(0)

                return FileResponse(buff, as_attachment=False, filename=f'boleta.pdf')


            elif documento_venta.tipodocumentoid == 1:
                if int(tipo_entrega) == 1:
                    Despacho.objects.create(
                        fechasolicitud = datetime.now().date(),
                        fechadespacho =  datetime.now().date(),
                        nroventa = ultima_ventas,
                        estadoid = Estado.objects.get(descripcion="Activo"),
                        tipodespacho = "Despacho"
                    )

                    ultimo_despacho = Despacho.objects.order_by('despachoid').last()
                    direccion_cliente = Direccioncliente.objects.get(clienteid=cliente_venta)
                    Guiadespacho.objects.create(
                        fechaguia = datetime.now().date(),
                        despachoid = ultimo_despacho,
                        iddircliente = direccion_cliente
                    )  


                else: 
                    Despacho.objects.create(
                        fechasolicitud = datetime.now().date(),
                        fechadespacho =  datetime.now().date(),
                        nroventa = ultima_ventas,
                        estadoid = Estado.objects.get(descripcion="Activo"),
                        tipodespacho = "Retiro")
                Boleta.objects.create(
                    fechaboleta = datetime.now().date(),
                    totalboleta = total_venta,
                    nroventa = ultima_ventas,
                    estadoid = Estado.objects.get(descripcion="Activo")
                )

                messages.warning(request, 'Venta realizada con exito')
                productos_venta.insert(0 , ["Nombre Producto", "Cantidad", "Total"]) 

                ultima_boleta = Boleta.objects.order_by('nroboleta').last()
                detalle_venta = Detalleventa.objects.filter(nroventa = ultima_boleta.nroventa)
                direccion_cliente = Direccioncliente.objects.get(clienteid=ultima_ventas.clienteid)
                if ultima_ventas.clienteid.personaid != None:
                    giro = "persona natural"
                else:
                    giro = "Distribuidor Ferreteria"

                response = HttpResponse(content_type=f'application/pdf')  
                buff = BytesIO()  
                c = canvas.Canvas(buff, pagesize=letter)
                c= generar_factura(c, ultima_ventas, ultima_boleta, detalle_venta, direccion_cliente, giro, 0)
                c.showPage()
                c.save()

                response.write(buff.getvalue())   
                buff.seek(0)

                return FileResponse(buff, as_attachment=False, filename=f'boleta.pdf')

        else:
            messages.warning(request, 'Ocurrio un error en la venta')

    return render(request, 'ventas/crear_venta.html',{'productos':productos, 'form':form, 'form_doc':form_doc, 'form_docu': form_docu, 'tipo_usuario': tipo_usuario})

@login_required(login_url="ingreso")
def listar_ventas(request):
    Seguimiento_paginas("Modulo Ventas - Listar Ventas", request.user)
    
    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    ventas = Venta.objects.all().order_by('nroventa').reverse

    if request.method == 'POST':
        if request.POST.get('VerVenta') is not None:
            request.session['_old_post'] = request.POST
            return HttpResponseRedirect('ver_venta')

    context = {
        'ventas': ventas,
        'tipo_usuario': tipo_usuario
    }

    return render(request, 'ventas/listar_ventas.html', context)

@login_required(login_url="ingreso")
def ver_venta(request):
    Seguimiento_paginas("Modulo Ventas - Ver Detalle Venta", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    old_post = request.session.get('_old_post')
    detalle_venta = Detalleventa.objects.filter(nroventa=old_post['VerVenta'])
    venta = Venta.objects.get(nroventa=old_post['VerVenta'])
    try:
        despacho = Despacho.objects.get(nroventa=old_post['VerVenta'])
        guia = Guiadespacho.objects.filter(despachoid=despacho)
    except Despacho.DoesNotExist:
        despacho= []
        guia= []
        
    context = {
        'detalle_venta': detalle_venta,
        'venta': venta,
        'despacho':guia,
        'tipo_usuario': tipo_usuario
    }

    return render(request, 'ventas/ver_venta.html', context)

#************************************Facturas**************************************

@login_required(login_url="ingreso")
def listar_facturas(request):
    Seguimiento_paginas("Modulo Facturas - Listar Facturas", request.user)
    
    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    facturas = Factura.objects.all()

    if request.method == 'POST':
        if request.POST.get('VerFactura') is not None:
            request.session['_old_post'] = request.POST
            return HttpResponseRedirect('ver_factura')
        elif request.POST.get('DescargarFactura') is not None:
            request.session['_old_post'] = request.POST
            return HttpResponseRedirect('ver_factura')

    context = {
        'facturas': facturas,
        'tipo_usuario': tipo_usuario
    }

    return render(request, 'facturas/listar_facturas.html', context)

@login_required(login_url="ingreso")
def ver_factura(request):
    Seguimiento_paginas("Modulo Facturas - Ver Detalle Factura", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    old_post = request.session.get('_old_post')

    factura = Factura.objects.get(numerofactura=old_post['VerFactura'])
    venta = Venta.objects.get(nroventa = factura.nroventa.nroventa)
    detalle_venta = Detalleventa.objects.filter(nroventa = factura.nroventa)
    direccion_cliente = Direccioncliente.objects.get(clienteid=factura.nroventa.clienteid)
    if venta.clienteid.personaid != None:
        giro = "persona natural"
    else:
        giro = "Distribuidor Ferreteria"

    if request.method == 'POST':

        tipoInforme = request.POST.get('informeCheck')
        descargarInforme = request.POST.get('descargarInforme')
        
        if tipoInforme is not None and descargarInforme is not None:

            lista1 = []
            lista2 = []
            lista3 = []

            lista1.append(["Nro Factura","Fecha","Neto","IVA","Total","Estado"])  
            val = Factura.objects.filter(numerofactura = old_post['VerFactura']).values_list('numerofactura','fechafactura','neto','iva','totalfactura','estadoid__descripcion')
            for valores in val:
                lista1.append(list(valores)) 

            lista2.append(["Nro Venta","Tipo pago","Run cliente","DV"])  
            val = Factura.objects.filter(numerofactura = old_post['VerFactura']).values_list('nroventa__nroventa','nroventa__tipodocumentoid__descripcion','nroventa__clienteid__personaid__runcuerpo','nroventa__clienteid__personaid__dv')
            for valores in val:
                lista2.append(list(valores))

            lista3.append(["Producto","Cantidad","Subtotal"])  
            val = Detalleventa.objects.filter(nroventa = factura.nroventa.nroventa).values_list('productoid__nombre','cantidad','subtotal')
            for valores in val:
                lista3.append(list(valores))

            if tipoInforme == "informeExcel":

                nombre_archivo = "Detalle Factura"
                tipo_doc = 'ms-excel'
                extension = 'xlsx'
                
                lista1.append("")
                lista2.append("")

                lista = lista1 + lista2 + lista3
                
                return  creacion_excel(nombre_archivo, lista, tipo_doc, extension)

            if tipoInforme == "informePdf":
                response = HttpResponse(content_type=f'application/pdf')  
                buff = BytesIO()  

                c = canvas.Canvas(buff, pagesize=letter)
                c= generar_factura(c, venta, factura, detalle_venta, direccion_cliente, giro, 1)
                c.showPage()
                c.save()

                response.write(buff.getvalue())   
                buff.seek(0)

                return FileResponse(buff, as_attachment=False, filename=f'factura.pdf')

            if tipoInforme == "informeWord":

                nombre_archivo = "Detalle Factura"
                # return  creacion_doc(lista1, nombre_archivo)
                document = Document()
                document.add_heading('Detalle Factura', 0)

                filas = 0
                for x in lista1:
                    columnas = len(x)
                    filas += 1

                # add grid table
                table = document.add_table(rows=filas, cols=columnas, style="Table Grid")

                for x in range(columnas):
                    table.rows[0].cells[x]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w'))))
                
                # access first row's cells
                heading_row = table.rows[0].cells

                # add headings
                cont = 0
                for value in lista1[0]:
                    heading_row[cont].text = value
                    cont += 1

                lista1.pop(0)
                cont = 0
                cont2 = 0

                for value in lista1:
                    cont += 1
                    data_row = table.rows[cont].cells

                    for x in value:
                        data_row[cont2].text = f'{x}'
                        cont2 += 1
                    cont2 = 0

                document.add_paragraph("")

                # parrafo.add_run().add_break()
                filas = 0
                for x in lista2:
                    columnas = len(x)
                    filas += 1

                # add grid table
                table2 = document.add_table(rows=filas, cols=columnas, style="Table Grid")

                for x in range(columnas):
                    table2.rows[0].cells[x]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w'))))

                # access first row's cells
                heading_row = table2.rows[0].cells

                # add headings
                cont = 0
                for value in lista2[0]:
                    heading_row[cont].text = value
                    cont += 1

                lista2.pop(0)
                cont = 0
                cont2 = 0

                for value in lista2:
                    cont += 1
                    data_row = table2.rows[cont].cells

                    for x in value:
                        data_row[cont2].text = f'{x}'
                        cont2 += 1
                    cont2 = 0 

                document.add_paragraph("")

                filas = 0
                for x in lista3:
                    columnas = len(x)
                    filas += 1

                # add grid table
                table3 = document.add_table(rows=filas, cols=columnas, style="Table Grid")

                for x in range(columnas):
                    table3.rows[0].cells[x]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w'))))

                # access first row's cells
                heading_row = table3.rows[0].cells

                # add headings
                cont = 0
                for value in lista3[0]:
                    heading_row[cont].text = value
                    cont += 1

                lista3.pop(0)
                cont = 0
                cont2 = 0

                for value in lista3:
                    cont += 1
                    data_row = table3.rows[cont].cells

                    for x in value:
                        data_row[cont2].text = f'{x}'
                        cont2 += 1
                    cont2 = 0  

                response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename={nombre_archivo}.docx'
                document.save(response)

                return response


    context = {
        'factura': factura,
        'tipo_usuario': tipo_usuario,
        'productos_boleta': detalle_venta
    }

    return render(request, 'facturas/ver_factura.html', context)

#************************************Notas de credito**************************************

@login_required(login_url="ingreso")
def Listar_notas_credito(request):
    Seguimiento_paginas("Modulo Notas de Credito - Listar Notas de Credito", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    notas_credito = Notacredito.objects.all()

    if request.method == 'POST':

        if request.POST.get('VerNotaCredito') is not None:
            request.session['_ver_nota_credito'] = request.POST
            return HttpResponseRedirect('ver_nota_credito')

    context = {
        'notas_credito': notas_credito,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'notas de credito/listar_notas_credito.html', context)

@login_required(login_url="ingreso")
def Ver_nota_credito(request):
    Seguimiento_paginas("Modulo Notas de Credito - Ver Detalle Nota de Credito", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    id_nota_credito = request.session.get('_ver_nota_credito')
    nota_credito = Notacredito.objects.get(nronota = id_nota_credito['VerNotaCredito'])

    lista = []
    lista.append(["Numero","Fecha","Total","Estado"])    
    val = Notacredito.objects.filter(nronota = id_nota_credito['VerNotaCredito']).values_list('nronota','fechanota','total','estadoid__descripcion')
    
    if request.method == 'POST':

        if request.POST.get('CambiarEstado') is not None:
            id_nota_credito = request.POST.get('CambiarEstado')
            Cambiar_estado_nota_credito(id_nota_credito)
            notacredito = Notacredito.objects.get(nronota=id_nota_credito)
            sweetify.success(request,
                f'La nota de credito {id_nota_credito} ha quedado {notacredito.estadoid.descripcion} correctamente')
            return redirect('listar_notas_credito')

        tipoInforme = request.POST.get('informeCheck')
        descargarInforme = request.POST.get('descargarInforme')
        
        if tipoInforme is not None and descargarInforme is not None:

            for valores in val:

                lista.append(list(valores))

            if tipoInforme == "informeExcel":

                nombre_archivo = "Compras"
                tipo_doc = 'ms-excel'
                extension = 'xlsx'
                
                return  creacion_excel(nombre_archivo, lista, tipo_doc, extension)

            if tipoInforme == "informePdf":

                response = HttpResponse(content_type=f'application/pdf')  
                buff = BytesIO()  

                if nota_credito.nroboleta != None:
                    documento_venta = Boleta.objects.get(nroboleta=nota_credito.nroboleta)
                else:
                    documento_venta = Factura.objects.get(numerofactura=nota_credito.numerofactura)

                venta = Venta.objects.get(nroventa = documento_venta.nroventa.nroventa)
                if venta.clienteid.personaid != None:
                    giro = "persona Natural"
                else:
                    giro = "Distribuidor Ferreteria"

                detalle_venta = []
                direccion_cliente = Direccioncliente.objects.get(clienteid=documento_venta.nroventa.clienteid)


                c = canvas.Canvas(buff, pagesize=letter)
                c= generar_factura(c, venta, nota_credito, detalle_venta, direccion_cliente, giro, 3)
                c.showPage()
                c.save()

                response.write(buff.getvalue())   
                buff.seek(0)

                return FileResponse(buff, as_attachment=False, filename=f'nota_credito.pdf')


            if tipoInforme == "informeWord":

                nombre_archivo = "Compras"
                return  creacion_doc(lista, nombre_archivo)

    context = {
        'tipo_usuario': tipo_usuario,
        'nota_credito': nota_credito
    }

    return render(request, 'notas de credito/ver_nota_credito.html', context)

def Cambiar_estado_nota_credito(id_nota_credito):
    notacredito = Notacredito.objects.get(nronota=id_nota_credito)

    if notacredito.estadoid.descripcion == 'Activo':
        notacredito.estadoid = Estado.objects.get(descripcion="Inactivo")
        notacredito.save()
    else:
        notacredito.estadoid = Estado.objects.get(descripcion="Activo")
        notacredito.save()

@login_required(login_url="ingreso")
def Seleccion_documento(request):

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    if request.method == 'POST':

        if request.POST.get('boleta') is not None:
            request.session['_documento'] = request.POST
            return HttpResponseRedirect('listar_documentos')

        else:
            request.session['_documento'] = request.POST
            return HttpResponseRedirect('listar_documentos')

    context = {
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'notas de credito/seleccion_documento.html', context)

@login_required(login_url="ingreso")
def Listar_documentos(request):

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    documento = request.session.get('_documento')

    if documento['documento'] == 'boleta':
        boletas = Boleta.objects.filter(estadoid__descripcion = 'Activo')
    else:
        boletas = None

    if documento['documento'] == 'factura':
        facturas = Factura.objects.filter(estadoid__descripcion = 'Activo')
    else:
        facturas = None

    if request.method == 'POST':

        if request.POST.get('CrearNotaCredito') is not None:
            request.session['_crear_nota_credito'] = request.POST
            return HttpResponseRedirect('crear_nota_credito')

    context = {
        'boletas': boletas,
        'facturas': facturas,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'notas de credito/listar_documentos.html', context)

@login_required(login_url="ingreso")
def Crear_nota_credito(request):
    Seguimiento_paginas("Modulo Notas de Credito - Crear Nota de Credito", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    documento = request.session.get('_crear_nota_credito')

    if documento['documento'] == 'boleta':
        detalles = Boleta.objects.filter(nroboleta = documento['id']).values('nroventa')
        detalles = Detalleventa.objects.filter(nroventa = detalles[0]['nroventa']).values('detalleventaid','productoid__nombre', 'productoid__precio')
    else:
        detalles = Factura.objects.filter(numerofactura = documento['id']).values('nroventa')
        detalles = Detalleventa.objects.filter(nroventa = detalles[0]['nroventa']).values('detalleventaid','productoid__nombre', 'productoid__precio')

    if request.method == 'POST':

        cont = 0
        suma2 = 0
        arreglo = []
        sumaPrecios = 0
        for key,value in request.POST.items():
            arreglo.append({key:value})

            cont += 1
            if cont > 2:
                suma2 += int(key)

        if arreglo[1]['TipoNotaCredito'] == 'completa':
            suma = 0 

            for x in detalles:
                suma += x['productoid__precio']

            
            now = datetime.now()

            if documento['documento'] == 'boleta':
                
                Notacredito.objects.create(
                    fechanota = now,
                    total = suma,
                    estadoid = Estado.objects.get(descripcion='Activo'),
                    nroboleta = documento['id']
                )
                sweetify.success(request, "Nota de credito creado correctamente")
                return redirect('listar_notas_credito')
            else:
                
                Notacredito.objects.create(
                    fechanota = now,
                    total = suma,
                    estadoid = Estado.objects.get(descripcion='Activo'),
                    numerofactura = Factura.objects.get(numerofactura = documento['id'])
                )
                sweetify.success(request, "Nota de credito creado correctamente")
                return redirect('listar_notas_credito')

        if arreglo[1]['TipoNotaCredito'] == 'parcial':
            
            from datetime import datetime
            now = datetime.now()

            if documento['documento'] == 'boleta':
                
                Notacredito.objects.create(
                    fechanota = now,
                    total = suma2,
                    estadoid = Estado.objects.get(descripcion='Activo'),
                    nroboleta = documento['id']
                )
                sweetify.success(request, "Nota de credito creado correctamente")
                return redirect('listar_notas_credito')
            else:
                
                Notacredito.objects.create(
                    fechanota = now,
                    total = suma2,
                    estadoid = Estado.objects.get(descripcion='Activo'),
                    numerofactura = Factura.objects.get(numerofactura = documento['id']) 
                )
                sweetify.success(request, "Nota de credito creada correctamente")
                return redirect('listar_notas_credito')
    else:
        sweetify.warning(request, "No es posible crear la nota de credito en este momento")

    context = {
        'detalles': detalles,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'notas de credito/crear_nota_credito.html', context)

#************************************Guias de despacho**************************************

@login_required(login_url="ingreso")
def listar_guias_despacho(request):

    Seguimiento_paginas("Modulo Guias de Despacho - Lista Guias de Despacho", request.user)
    
    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    guias_despacho = Guiadespacho.objects.all()

    if request.method == 'POST':
        if request.POST.get('VerGuiaDespacho') is not None:
            request.session['_old_post'] = request.POST
            return HttpResponseRedirect('ver_guia_despacho')

    context = {
        'guias_despacho': guias_despacho,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'guias_despacho/listar_guias_despacho.html', context)

@login_required(login_url="ingreso")
def ver_guia_despacho(request):

    Seguimiento_paginas("Modulo Guias de Despacho - Detalle Guia de Despacho", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    old_post = request.session.get('_old_post')

    guia = Guiadespacho.objects.get(nroguia=old_post['VerGuiaDespacho'])

    venta = Venta.objects.get(nroventa = guia.despachoid.nroventa.nroventa)
    detalle_venta = Detalleventa.objects.filter(nroventa = venta.nroventa)
    direccion_cliente = Direccioncliente.objects.get(iddircliente=guia.iddircliente.iddircliente)
    if venta.clienteid.personaid != None:
        giro = "persona natural"
    else:
        giro = "Distribuidor Ferreteria"

    if request.method == 'POST':

        tipoInforme = request.POST.get('informeCheck')
        descargarInforme = request.POST.get('descargarInforme')
        
        if tipoInforme is not None and descargarInforme is not None:

            lista1 = []
            lista2 = []

            lista1.append(["Folio Documento","Fecha","Direccion"])  
            val = Guiadespacho.objects.filter(nroguia = old_post['VerGuiaDespacho']).values_list('nroguia','fechaguia','iddircliente__direccionid')

            for valores in val:
                lista1.append(list(valores)) 
            
            lista2.append(["Nro Venta","Tipo pago","Run cliente","DV"])  
            val = Guiadespacho.objects.filter(nroguia = old_post['VerGuiaDespacho']).values_list('despachoid__nroventa__nroventa','despachoid__nroventa__tipodocumentoid__descripcion','despachoid__nroventa__clienteid__personaid__runcuerpo','despachoid__nroventa__clienteid__personaid__dv')
            for valores in val:
                lista2.append(list(valores))
                
            if tipoInforme == "informeExcel":

                nombre_archivo = "Detalle Guia Despacho"
                tipo_doc = 'ms-excel'
                extension = 'xlsx'
                
                lista1.append("")
                lista2.append("")

                lista = lista1 + lista2 
                
                return  creacion_excel(nombre_archivo, lista, tipo_doc, extension)

            if tipoInforme == "informePdf":
                response = HttpResponse(content_type=f'application/pdf')  
                buff = BytesIO()  

                c = canvas.Canvas(buff, pagesize=letter)
                c= generar_factura(c, venta, guia, detalle_venta, direccion_cliente, giro, 2)
                c.showPage()
                c.save()

                response.write(buff.getvalue())   
                buff.seek(0)

                return FileResponse(buff, as_attachment=False, filename=f'guia_despacho.pdf')

            if tipoInforme == "informeWord":

                nombre_archivo = "Detalle Guia Despacho"
                document = Document()
                document.add_heading('Detalle Guia Despacho', 0)

                filas = 0
                for x in lista1:
                    columnas = len(x)
                    filas += 1

                # add grid table
                table = document.add_table(rows=filas, cols=columnas, style="Table Grid")

                for x in range(columnas):
                    table.rows[0].cells[x]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w'))))
                
                # access first row's cells
                heading_row = table.rows[0].cells

                # add headings
                cont = 0
                for value in lista1[0]:
                    heading_row[cont].text = value
                    cont += 1

                lista1.pop(0)
                cont = 0
                cont2 = 0

                for value in lista1:
                    cont += 1
                    data_row = table.rows[cont].cells

                    for x in value:
                        data_row[cont2].text = f'{x}'
                        cont2 += 1
                    cont2 = 0

                document.add_paragraph("")

                # parrafo.add_run().add_break()
                filas = 0
                for x in lista2:
                    columnas = len(x)
                    filas += 1

                # add grid table
                table2 = document.add_table(rows=filas, cols=columnas, style="Table Grid")

                for x in range(columnas):
                    table2.rows[0].cells[x]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w'))))

                # access first row's cells
                heading_row = table2.rows[0].cells

                # add headings
                cont = 0
                for value in lista2[0]:
                    heading_row[cont].text = value
                    cont += 1

                lista2.pop(0)
                cont = 0
                cont2 = 0

                for value in lista2:
                    cont += 1
                    data_row = table2.rows[cont].cells

                    for x in value:
                        data_row[cont2].text = f'{x}'
                        cont2 += 1
                    cont2 = 0 

                document.add_paragraph("")

                filas = 0

                response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename={nombre_archivo}.docx'
                document.save(response)

                return response


    context = {
        'guia': guia,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'guias_despacho/ver_guia_despacho.html', context)

#************************************Informes**************************************

@login_required(login_url="ingreso")
def informe_productos(request):

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    Seguimiento_paginas("Informe Productos", request.user)
    
    form1 = addproductsForm(request.POST)

    if request.method == 'POST':
        vistaPrevia = request.POST.get('vistaPrevia')
        descargarInforme = request.POST.get('descargarInforme')
        productos = request.POST.get('productos')
        precio = request.POST.get('precio')
        nombre = request.POST.get('nombre')
        precioCompra = request.POST.get('precioCompra')
        stockCritico = request.POST.get('stockCritico')
        fechaVencimiento = request.POST.get('fechaVencimiento')
        codigoBarra = request.POST.get('codigoBarra')
        stock_producto = request.POST.get('stock')
        stockCheck = request.POST.get('stockCheck')
        estado = request.POST.get('estado')
        estadoCheck = request.POST.get('estadoCheck')
        familiaProducto = request.POST.get('familiaProducto')
        familiaProductoCheck = request.POST.get('familiaProductoCheck')
        nomFamiliaProducto = request.POST.get('familia_producto')
        tipoProducto = request.POST.get('tipoProducto')
        nomTipoProducto = request.POST.get('tipo_producto')
        tipoProductoCheck = request.POST.get('tipoProductoCheck')
        familiaproid = request.POST.get('familiaproid')
        tipoproductoid = request.POST.get('tipoproductoid')

        lista = []
        visitas = []

        tipoInforme = request.POST.get('informeCheck')
        
        np_array = []
        np_array = np.array(lista)

        if productos == "on":
            productos = Producto.objects.filter().values_list("nombre","precio","stock","stockcritico","fechavencimiento","codigo","familiaproid__descripcion","tipoproductoid__descripcion","estadoid__descripcion","bodegaid__pasillo","bodegaid__estante","bodegaid__casillero").order_by("productoid")
            if stockCheck == "conStock":
                productos = productos.filter(stock__gt=0)
            if stockCheck == "sinStock":
                productos = productos.filter(stock__lt=0)
            if estadoCheck == "disponible":
                productos = productos.filter(estadoid__estadoid=1)
            if estadoCheck == "noDisponible":
                productos = productos.filter(estadoid__estadoid=2)
            if familiaProductoCheck == "porNombreF":
                productos = productos.filter(familiaproid__familiaproid = familiaproid)
            if tipoProductoCheck == "porNombreT":
                productos = productos.filter(tipoproductoid__tipoproductoid = tipoproductoid)
            columnas = (["Nombre","Precio","Stock","Stock Critico","Fecha Vencimiento","Codigo","Familia Producto","Tipo Producto","Estado","Pasillo","Estante","Casillero"])
 
            np_array = np.array(productos)
            df = pd.DataFrame(np_array, columns = columnas)
            df["Bodega"] = "Pasillo:" +df["Pasillo"] + " Estante:" + df["Estante"]+ " Casillero:" + df["Casillero"]
            df = df.drop(['Pasillo'], axis=1)
            df = df.drop(['Estante'], axis=1)
            df = df.drop(['Casillero'], axis=1)

            if nombre == None:
                df = df.drop(['Nombre'], axis=1)
            if precio == None:
                df = df.drop(['Precio'], axis=1)
            if stock_producto == None:
                df = df.drop(['Stock'], axis=1)
            if stockCritico == None:
                df = df.drop(['Stock Critico'], axis=1)
            if estado == None:
                df = df.drop(['Estado'], axis=1)
            if fechaVencimiento == None:
                df = df.drop(['Fecha Vencimiento'], axis=1)
            if codigoBarra == None:
                df = df.drop(['Codigo'], axis=1)
            if codigoBarra == None:
                df = df.drop(['Bodega'], axis=1)
            if familiaProducto == None:
                df = df.drop(['Familia Producto'], axis=1)
            if tipoProducto == None:
                df = df.drop(['Tipo Producto'], axis=1)

            if df.shape[1] > 6:
                length_dataframe = df.shape[1]
                df2= df.iloc[:, 6:int(length_dataframe)] 
                df= df.iloc[:, 0:6]

                lista_productos2 = df2.values.tolist() 
                columnas_df2 = df2.columns.values.tolist() 
                lista_productos2.insert(0, columnas_df2)

                lista_productos = df.values.tolist() 
                columnas_df = df.columns.values.tolist() 
                lista_productos.insert(0, columnas_df)
            else:
                lista_productos = df.values.tolist() 
                columnas_df = df.columns.values.tolist() 
                lista_productos.insert(0, columnas_df)

            if tipoInforme == "informeExcel":
                
                nombre_archivo = "Productos"
                tipo_doc = 'ms-excel'
                extension = 'xlsx'
                return  creacion_excel(nombre_archivo, lista, tipo_doc, extension)
                

            if tipoInforme == "informePdf":

                tipo_doc = 'pdf'
                extension = 'pdf'
                nombre = 'Informe Productos'
                if vistaPrevia:
                    if df.shape[1] > 6:
                        return creacion_pdf(lista_productos,tipo_doc,A4,nombre,extension, lista_productos2, valor=False)
                    else:
                        return creacion_pdf(lista_productos,tipo_doc,A4,nombre,extension,valor=False)
                else:
                    if df.shape[1] > 6:
                        return creacion_pdf(lista_productos,tipo_doc,A4,nombre,extension, lista_productos2, valor=True)
                    else:
                        return creacion_pdf(lista_productos,tipo_doc,A4,nombre,extension,valor=True)
            if tipoInforme == "informeWord": 

                tipo_doc = 'ms-word'
                extension = 'docx'
                
                nombre = 'Productos'
                if vistaPrevia:
                    if visitas == []:
                        
                        return creacion_doc(lista_productos,nombre)
                    else:
                        return creacion_doc(lista_productos,nombre)
                else:
                    if visitas == []:
                        return creacion_doc(lista_productos,nombre)
                    else:
                        return creacion_doc(lista_productos,nombre)

    context = {
        'form1': form1,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'informes/informe_productos.html', context)

@login_required(login_url="ingreso")
def informe_proveedores(request):

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None
    
    form1 = FormProveedor(request.POST)
    Seguimiento_paginas("Informe Proveedores", request.user)

    if request.method == 'POST':
        vistaPrevia = request.POST.get('vistaPrevia')
        proveedores = request.POST.get('proveedores')
        correoP = request.POST.get('correoP')
        telefonoP = request.POST.get('telefonoP')
        direccionP = request.POST.get('direccionP')
        estadoP = request.POST.get('estadoP')
        estadoProveedorCheck = request.POST.get('estadoProveedorCheck')
        nomProveedor = request.POST.get('proveedor')
        categoriaP = request.POST.get('categoriaP')
        categoriaPCheck = request.POST.get('categoriaPCheck')
        nomCategoria = request.POST.get('categoria_proveedor')
        rubroid = request.POST.get('rubroid')

        lista = []
        visitas = []

        tipoInforme = request.POST.get('informeCheck')
        
        np_array = []
        np_array = np.array(lista)

        if proveedores == "on":
            proveedores = Proveedor.objects.all().values_list("razonsocial","rutcuerpo","dv","fono","rubroid__descripcion","direccionid__calle","direccionid__numero","direccionid__comunaid__nombre","estadoid__descripcion").order_by("proveedorid")
            columnas = (["Razon Social","RUN", "DV", "Telefono", "Rubro","calle","numero","comuna","Estado"])

            if estadoProveedorCheck == "activosP":
                proveedores = proveedores.filter(estadoid__estadoid=1)
            if estadoProveedorCheck == "bloqueadosP":
                proveedores = proveedores.filter(estadoid__estadoid=2)
            if categoriaPCheck == "porcategoriaP":
                proveedores = proveedores.filter(rubroid__rubroid = rubroid)
                
            np_array = np.array(proveedores)
            df = pd.DataFrame(np_array, columns = columnas)

            df["Direccion"] = df["calle"] + " " + df["numero"] + "," + df["comuna"]
            df = df.drop(['calle'], axis=1)
            df = df.drop(['numero'], axis=1)
            df = df.drop(['comuna'], axis=1)

            if telefonoP == None:
                df = df.drop(['Telefono'], axis=1)
            if direccionP == None:
                df = df.drop(['Direccion'], axis=1)

            if df.shape[1] > 6:
                length_dataframe = df.shape[1]
                df2 = df.iloc[:, 6:int(length_dataframe)] 
                df = df.iloc[:, 0:6]

                lista_proveedores2 = df2.values.tolist() 
                columnas_df2 = df2.columns.values.tolist() 
                lista_proveedores2.insert(0, columnas_df2)

                lista_proveedores = df.values.tolist() 
                columnas_df = df.columns.values.tolist() 
                lista_proveedores.insert(0, columnas_df)
            else:
                lista_proveedores = df.values.tolist() 
                columnas_df = df.columns.values.tolist() 
                lista_proveedores.insert(0, columnas_df)

            if tipoInforme == "informeExcel":
                
                nombre_archivo = "Proveedores"
                tipo_doc = 'ms-excel'
                extension = 'xlsx'
                creacion_excel(nombre_archivo, lista_proveedores, tipo_doc, extension)
                
                
            if tipoInforme == "informePdf":

                tipo_doc = 'pdf'
                extension = 'pdf'
                nombre = 'Informe Proveedores'
                if vistaPrevia:
                    if df.shape[1] >= 6:
                        return creacion_pdf(lista_proveedores,tipo_doc,A4,nombre,extension, lista_proveedores2, valor=False)
                    else:
                        return creacion_pdf(lista_proveedores,tipo_doc,A4,nombre,extension,valor=False)
                else:
                    if df.shape[1] >= 6:
                        return creacion_pdf(lista_proveedores,tipo_doc,A4,nombre,extension, lista_proveedores2, valor=True)
                    else:
                        return creacion_pdf(lista_proveedores,tipo_doc,A4,nombre,extension,valor=True)

            if tipoInforme == "informeWord": 

                tipo_doc = 'ms-word'
                extension = 'docx'
                
                nombre = 'Proveedores'
                if vistaPrevia:
                    if visitas == []:
                        return creacion_doc(lista_proveedores,nombre)
                    else:
                        return creacion_doc(lista_proveedores,nombre)
                else:
                    if visitas == []:
                        return creacion_doc(lista_proveedores,nombre)
                    else:
                        return creacion_doc(lista_proveedores,nombre)

    context = {
        'form1':form1,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'informes/informe_proveedores.html', context)

@login_required(login_url="ingreso")
def Seleccion_informe(request):

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None
    
    context = {
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'informes/seleccion_informe.html', context)

@login_required(login_url="ingreso")
def informe_pedidos(request):

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    Seguimiento_paginas("Informe Pedidos", request.user)

    form1 = FormOrdencompra(request.POST)

    if request.method == 'POST':
        ordenes = request.POST.get('ordenes')
        fechaOrdenP = request.POST.get('fechaOrdenP')
        fechaLlegadaOrdenP = request.POST.get('fechaLlegadaOrdenP')
        fechaRecepOrdenP = request.POST.get('fechaRecepOrdenP')
        horaRecepOrdenP = request.POST.get('horaRecepOrdenP')
        cantidadOrdenP = request.POST.get('cantidadOrdenP')
        productoOrdenP = request.POST.get('productoOrdenP')
        estadoordenid = request.POST.get('estadoordenid')
        proveedorid = request.POST.get('proveedorid')

        estadoRecepcionO = request.POST.get('estadoRecepcionO')
        razonSocialOrdenP = request.POST.get('razonSocialOrdenP')
        razonSocialOrdenPCheck = request.POST.get('razonSocialOrdenPCheck')
        nomProveedor = request.POST.get('proveedor')
        tipoInforme = request.POST.get('informeCheck')
        vistaPrevia = request.POST.get('vistaPrevia')
        visitasPagina = request.POST.get('visitas')
        
        lista = []
        visitas = []

        np_array = []
        np_array = np.array(lista)

        if ordenes == "on":
            ordenes = Detalleorden.objects.all().values_list("ordenid__proveedorid__razonsocial","ordenid__fechapedido","ordenid__estadoordenid__descripcion","productoid__nombre","cantidad").order_by("detalleid")
            columnas = (["Razon Social","Fecha Pedido", "Estado", "Producto", "Cantidad"])
            if estadoRecepcionO == "on":   
                ordenes = ordenes.filter(ordenid__estadoordenid = estadoordenid)
            if razonSocialOrdenPCheck == "porRazonSocialOrdenP":
                ordenes = ordenes.filter(ordenid__proveedorid = proveedorid)

            if len(ordenes) < 1:
                sweetify.warning(request, 'No existen registros con los par??metros seleccionados')
                return redirect('informe_pedidos')
            else:
                np_array = np.array(ordenes)
                df = pd.DataFrame(np_array, columns = columnas)

                if cantidadOrdenP == None:
                    df = df.drop(['Cantidad'], axis=1)
                if productoOrdenP == None:
                    df = df.drop(['Producto'], axis=1)

                lista_proveedores = df.values.tolist() 
                columnas_df = df.columns.values.tolist() 
                lista_proveedores.insert(0, columnas_df)
            
            if tipoInforme == "informeExcel":
                nombre_archivo = "Pedidos"
                tipo_doc = 'ms-excel'
                extension = 'xlsx'
                return  creacion_excel(nombre_archivo, lista_proveedores, tipo_doc, extension)
                

            if tipoInforme == "informePdf":
                tipo_doc = 'pdf'
                extension = 'pdf'
                nombre = 'Informe Pedidos'
                if vistaPrevia:
                    return creacion_pdf(lista_proveedores,tipo_doc,A4,nombre,extension, valor=False)
                else:
                    return creacion_pdf(lista_proveedores,tipo_doc,A4,nombre,extension, valor=True)

            if tipoInforme == "informeWord": 
                tipo_doc = 'ms-word'
                extension = 'docx'
                nombre = 'Pedidos'
                if vistaPrevia:
                    if visitas == []:
                        return creacion_doc(lista_proveedores,nombre)
                    else:
                        return creacion_doc(lista_proveedores,nombre)
                else:
                    if visitas == []:
                        return creacion_doc(lista_proveedores,nombre)
                    else:
                        return creacion_doc(lista_proveedores,nombre)

    context = {
        'form1':form1,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'informes/informe_pedidos.html', context)

@login_required(login_url="ingreso")
def informe_ventas(request):

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    form1 = FormVenta(request.POST)

    Seguimiento_paginas("Informe Ventas", request.user)
    
    if request.method == 'POST':

        ventas = request.POST.get('ventas')
        fechaVenta = request.POST.get('fechaVenta')
        montoTotal = request.POST.get('montoTotal')
        nombreCliente = request.POST.get('nombreCliente')
        nombreProducto = request.POST.get('nombreProducto')
        cantidadVenta = request.POST.get('cantidadVenta')
        subTotal = request.POST.get('subTotal')

        DocuTributario = request.POST.get('DocuTributario')
        docuTributarioCheck = request.POST.get('docuTributarioCheck')
        tipodocumentoid = request.POST.get('tipodocumentoid')
        tipoPago = request.POST.get('tipoPago')
        TipoPagoCheck = request.POST.get('TipoPagoCheck')
        tipopagoid = request.POST.get('tipopagoid')

        tipoInforme = request.POST.get('informeCheck')
        vistaPrevia = request.POST.get('vistaPrevia')
        
        lista = []
        visitas = []

        tipoInforme = request.POST.get('informeCheck')
        
        np_array = []
        np_array = np.array(lista)

        if ventas == "on":
            ventas = Detalleventa.objects.all().values_list("nroventa","nroventa__fechaventa","nroventa__totalventa","nroventa__clienteid__personaid__nombres","nroventa__clienteid__empresaid__razonsocial", "productoid__nombre","cantidad","subtotal","nroventa__tipodocumentoid__descripcion","nroventa__tipopagoid__descripcion").order_by("nroventa")

            columnas = (["Nro Venta","Fecha Venta", "Valor Total", "Cliente", "Empresa", "Nombre Producto", "Cantidad", "Sub Total", "Documento Venta", "Tipo de Pago"])
            
            if docuTributarioCheck == "porDocuTributario":   
                ventas = ventas.filter(nroventa__tipodocumentoid = tipodocumentoid)
            if TipoPagoCheck == "porTipoPago":
                ventas = ventas.filter(nroventa__tipopagoid = tipopagoid)

            if len(ventas) < 1:
                sweetify.warning(request, 'No existen registros con los par??metros seleccionados')
                return redirect('informe_ventas')
            else:
                np_array = np.array(ventas)
                df = pd.DataFrame(np_array, columns = columnas)

                df["Nombre Cliente"] = df["Cliente"] + df["Empresa"]
                df = df.drop(['Cliente'], axis=1)
                df = df.drop(['Empresa'], axis=1)

                if fechaVenta == None:
                    df = df.drop(['Fecha Venta'], axis=1)
                if montoTotal == None:
                    df = df.drop(['Valor Total'], axis=1)
                if nombreCliente == None:
                    df = df.drop(['Nombre Cliente'], axis=1)
                if nombreProducto == None:
                    df = df.drop(['Nombre Producto'], axis=1)
                if cantidadVenta == None:
                    df = df.drop(['Cantidad'], axis=1)
                if subTotal == None:
                    df = df.drop(['Sub Total'], axis=1)

                if DocuTributario == None:
                    df = df.drop(['Documento Venta'], axis=1)
                if tipoPago == None:
                    df = df.drop(['Tipo de Pago'], axis=1)

            if df.shape[1] > 5:
                length_dataframe = df.shape[1]
                df2= df.iloc[:, 6:int(length_dataframe)] 
                df= df.iloc[:, 0:6]

                lista_ventas2 = df2.values.tolist() 
                columnas_df2 = df2.columns.values.tolist() 
                lista_ventas2.insert(0, columnas_df2)

                lista_ventas = df.values.tolist() 
                columnas_df = df.columns.values.tolist() 
                lista_ventas.insert(0, columnas_df)
            else:
                lista_ventas = df.values.tolist() 
                columnas_df = df.columns.values.tolist() 
                lista_ventas.insert(0, columnas_df)

            if tipoInforme == "informeExcel":
                nombre_archivo = "Informe Detallado Ventas"
                tipo_doc = 'ms-excel'
                extension = 'xlsx'
                return  creacion_excel(nombre_archivo, lista_ventas, tipo_doc, extension)
                

            if tipoInforme == "informePdf":
                tipo_doc = 'pdf'
                extension = 'pdf'
                nombre = 'Informe Detallado Ventas'
                if vistaPrevia:
                    if df.shape[1] > 5:
                        return creacion_pdf(lista_ventas,tipo_doc,A4,nombre,extension, lista_ventas2, valor=False)
                    else:
                        return creacion_pdf(lista_ventas,tipo_doc,A4,nombre,extension,valor=False)
                else:
                    if df.shape[1] > 5:
                        return creacion_pdf(lista_ventas,tipo_doc,A4,nombre,extension, lista_ventas2, valor=True)
                    else:
                        return creacion_pdf(lista_ventas,tipo_doc,A4,nombre,extension,valor=True)

            if tipoInforme == "informeWord": 
                tipo_doc = 'ms-word'
                extension = 'docx'
                nombre = 'Informe Detallado Ventas'
                if vistaPrevia:
                    if visitas == []:
                        return creacion_doc(lista_ventas,nombre)
                    else:
                        return creacion_doc(lista_ventas,nombre)
                else:
                    if visitas == []:
                        return creacion_doc(lista_ventas,nombre)
                    else:
                        return creacion_doc(lista_ventas,nombre)

    context = {
        'form1':form1,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'informes/informe_ventas.html', context)

@login_required(login_url="ingreso")
def informe_visitas(request):

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    Seguimiento_paginas("Informe Visitas", request.user)

    form1 = FormAccionpagina(request.POST)

    if request.method == 'POST':

        visitas = request.POST.get('visitas')
        fechaIn = request.POST.get('fechaIn')
        nombreModulo = request.POST.get('nombreModulo')
        fechain = request.POST.get('fechain')
        fechaInicio = request.POST.get('fechaInicio')
        porModuloCheck = request.POST.get('porModuloCheck')
        moduloCheck = request.POST.get('moduloCheck')
        detalleUsuario = request.POST.get('detalleUsuario')
        nombreUsuario = request.POST.get('nombreUsuario')
        correoUsuario = request.POST.get('correoUsuario')
        rolUsuario = request.POST.get('rolUsuario')

        tipoInforme = request.POST.get('informeCheck')
        vistaPrevia = request.POST.get('vistaPrevia')
        
        lista = []

        np_array = []
        np_array = np.array(lista)

        if visitas == "on":
            visitas = Accionpagina.objects.all().values_list("modulo","fechain","usuarioid__empresaid__razonsocial","usuarioid__personaid__nombres","usuarioid__email","usuarioid__rolid__descripcion").order_by("-fechain")
            
            columnas = (["Modulo", "Fecha Registro", "Empresa", "Persona", "Correo","Cargo"])

            if porModuloCheck == "porModulo":   
                visitas = visitas.filter(modulo__icontains = moduloCheck)

            if fechaInicio == "on":
                visitas = visitas.filter(fechain__gte = fechain)

            if len(visitas) < 1:
                sweetify.warning(request, 'No existen registros con los par??metros seleccionados')
                return redirect('informe_visitas')
            else:

                np_array = np.array(visitas)
                df = pd.DataFrame(np_array, columns = columnas)

                df["Usuario"] = df["Empresa"] + df["Persona"] 
                df = df.drop(['Empresa'], axis=1)
                df = df.drop(['Persona'], axis=1)

                df['Fecha Registro'] = df['Fecha Registro'].astype(str).str.slice(0, 19)

                if fechaIn == None:
                    df = df.drop(['Fecha Registro'], axis=1)
                if nombreModulo == None:
                    df = df.drop(['Modulo'], axis=1)
                if nombreUsuario == None:
                    df = df.drop(['Usuario'], axis=1)
                if correoUsuario == None:
                    df = df.drop(['Correo'], axis=1)
                if rolUsuario == None:
                    df = df.drop(['Cargo'], axis=1)
                
                if df.shape[1] > 6:
                    length_dataframe = df.shape[1]
                    df2 = df.iloc[:, 6:int(length_dataframe)] 
                    df = df.iloc[:, 0:6]

                    lista_visitas2 = df2.values.tolist() 
                    columnas_df2 = df2.columns.values.tolist() 
                    lista_visitas2.insert(0, columnas_df2)

                    lista_visitas = df.values.tolist() 
                    columnas_df = df.columns.values.tolist() 
                    lista_visitas.insert(0, columnas_df)
                else:
                    lista_visitas = df.values.tolist() 
                    columnas_df = df.columns.values.tolist() 
                    lista_visitas.insert(0, columnas_df)
                
            if tipoInforme == "informeExcel":
                nombre_archivo = "Pedidos"
                tipo_doc = 'ms-excel'
                extension = 'xlsx'
                return  creacion_excel(nombre_archivo, lista_visitas, tipo_doc, extension)
                

            if tipoInforme == "informePdf":
                tipo_doc = 'pdf'
                extension = 'pdf'
                nombre = 'Informe Pedidos'
                if vistaPrevia:
                    if df.shape[1] > 6:
                        return creacion_pdf(lista_visitas,tipo_doc,A4,nombre,extension, lista_visitas2, valor=False)
                    else:
                        return creacion_pdf(lista_visitas,tipo_doc,A4,nombre,extension,valor=False)
                else:
                    if df.shape[1] > 6:
                        return creacion_pdf(lista_visitas,tipo_doc,A4,nombre,extension, lista_visitas2, valor=True)
                    else:
                        return creacion_pdf(lista_visitas,tipo_doc,A4,nombre,extension,valor=True)

            if tipoInforme == "informeWord": 
                tipo_doc = 'ms-word'
                extension = 'docx'
                nombre = 'Pedidos'
                if vistaPrevia:
                    if visitas == []:
                        return creacion_doc(lista_visitas,nombre)
                    else:
                        return creacion_doc(lista_visitas,nombre)
                else:
                    if visitas == []:
                        return creacion_doc(lista_visitas,nombre)
                    else:
                        return creacion_doc(lista_visitas,nombre)

    context = {
        'form1':form1,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'informes/informe_visitas.html', context)
  
#************************************Compras cliente**************************************

def Comprar(request):

    # Seguimiento_paginas("Compras - Cliente", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    productos = Producto.objects.filter(estadoid__descripcion = 'Activo')

    context = {
        'tipo_usuario': tipo_usuario,
        'productos': productos
    }

    return render(request, 'compras/comprar.html', context)

def Procesar_compra(request):

    # Seguimiento_paginas("Procesar compra - Cliente", request.user)

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    if request.user:
        if Usuario.objects.filter(nombreusuario=request.user).values('empresaid'):
            cliente_empresa_existe = 'si'
        else:
            cliente_empresa_existe = None

    if request.method == 'POST': 
        cont = 0
        cont2 = 0
        detalle = []
        venta = []
        correo = []
        tipo_documento = []

        for key,value in request.POST.items():
            cont2 += 1

        for key,value in request.POST.items():

            if cont > 0 and cont < (cont2 - 4): 
                detalle.append({key:value})
            
            if cont >= (cont2 - 4) and cont < (cont2 - 2): 
                venta.append({key:value})

            if cont == (cont2 -2): 
                tipo_documento.append({key:value})
            
            if cont == (cont2 -1): 
                correo.append({key:value})
                
            cont += 1

        now = datetime.now()
        doc = ''
        giro = ''
        if tipo_documento[0]['documentoCheck'] == 'boleta':
            documento = Tipodocumento.objects.get(descripcion = 'Boleta')
            doc = 0
            giro = 'Persona natural'
        else:
            documento = Tipodocumento.objects.get(descripcion = 'Factura')
            doc = 1
            giro = 'Persona natural'

        tipo_pago = Tipopago.objects.get(descripcion = 'Debito')

        total = (venta[1]['total_peso']).translate(str.maketrans('', '', string.punctuation))
        Venta.objects.create(
            fechaventa = now,
            totalventa = int(total),
            tipodocumentoid = documento,
            clienteid = Cliente.objects.get(clienteid = 266),
            tipopagoid = tipo_pago
        )

        cont = 0
        cont2 = 0

        venta_id = Venta.objects.filter(
            fechaventa = now,
            totalventa = int(total),
            tipodocumentoid = documento,
            tipopagoid = tipo_pago
        ).values('nroventa').last()

        for value in detalle:
            cont += 1
            
            if cont == 4:
                cont2 += 1
                cont= 0
        
        cont_valores = 0
        for x in range(cont2):
            
            if x > 0:
                cont_valores += 4

            Detalleventa.objects.create(
                cantidad = int(detalle[3 + cont_valores][f'producto_cantidad{x}']),
                subtotal = int((detalle[2 + cont_valores][f'producto_precio{x}'])[1:]),
                productoid = Producto.objects.get(productoid = int(detalle[0 + cont_valores][f'producto_id{x}'])),
                nroventa = Venta.objects.get(nroventa = venta_id['nroventa'])
            )
            producto_modificar = Producto.objects.get(productoid = int(detalle[0 + cont_valores][f'producto_id{x}']))
            producto_vendido, created =  Producto.objects.get_or_create(productoid = producto_modificar.productoid)
            producto_vendido.stock = producto_vendido.stock-int(detalle[3 + cont_valores][f'producto_cantidad{x}'])
            producto_vendido.save()
            cont += 1

        doc_adjunto = ''

        if tipo_documento[0]['documentoCheck'] == 'boleta':
            Boleta.objects.create(
                fechaboleta = now,
                totalboleta = int(total),
                nroventa = Venta.objects.get(nroventa = venta_id['nroventa']),
                estadoid = Estado.objects.get(descripcion = 'Activo')
            )
            doc_adjunto = Boleta.objects.get(nroventa = venta_id['nroventa'])
        else:
            Factura.objects.create(
                fechafactura = now,
                neto = int(total)-(int(total)*0.19),
                iva = int(total)*0.19,
                totalfactura = int(total),
                nroventa = Venta.objects.get(nroventa = venta_id['nroventa']),
                estadoid = Estado.objects.get(descripcion = 'Activo')
            )
            doc_adjunto = Factura.objects.get(nroventa = venta_id['nroventa'])

        
        venta_correo = Venta.objects.filter(nroventa = venta_id['nroventa']).values_list('nroventa', 'fechaventa', 'totalventa')
        detalle_venta_correo = Detalleventa.objects.filter(nroventa = venta_id['nroventa']).values_list('productoid__nombre', 'cantidad', 'subtotal')

        message = loader.render_to_string(
            'compras/enviar_detalle.html',
            {
                'venta': venta_correo,
                'detalle_venta': detalle_venta_correo
            }
        )

        de_email = settings.EMAIL_HOST_USER
        para_email = correo[0]['correo']
        # mime_message = MIMEMultipart()
        mime_message = MIMEText(message, "html", _charset="utf-8")
        mime_message["From"] = de_email
        mime_message["To"] = para_email
        mime_message["Subject"] = "Detalle de la compra realizada"

        response = HttpResponse(content_type=f'application/pdf') 
        contenido = "attachment; filename = {0}.{1}".format('Documento', '.pdf')
        response["Content-Disposition"] = contenido 

        buff = BytesIO()  


        c = canvas.Canvas(buff, pagesize=letter)

        venta_adjunto = Venta.objects.get(nroventa = venta_id['nroventa'])
        
        productos_adjunto = Detalleventa.objects.filter(nroventa = venta_id['nroventa'])
        direccion_cliente = Direccioncliente.objects.get(clienteid=266)

        c= generar_factura(c, venta_adjunto, doc_adjunto, productos_adjunto, direccion_cliente, giro, doc)
        c.showPage()
        c.save()

        response.write(buff.getvalue())   
        buff.seek(0)

        smtpObj = smtplib.SMTP(settings.EMAIL_HOST, 587)
        smtpObj.login(de_email, settings.EMAIL_HOST_PASSWORD)
        smtpObj.sendmail(de_email, para_email, mime_message.as_string())
        
        sweetify.success(request, "Venta realizada con exito")

        return response
        
    context = {
        'tipo_usuario': tipo_usuario,
        'cliente_empresa_existe': cliente_empresa_existe
    }
        

    return render(request, 'compras/procesar_compra.html', context)

def dashboard(request):
    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    # 1 Boleta
    # 2 Factura
    ventas = Venta.objects.all()

    boletas = Venta.objects.filter(tipodocumentoid__tipodocumentoid = 1)
    facturas = Venta.objects.filter(tipodocumentoid__tipodocumentoid = 2)
    ventas_tipodocumento = [len(boletas), len(facturas)]

    despachos = Despacho.objects.all()
    
    total_ventas_despachos = [len(ventas)-len(despachos), len(despachos)]

    detalle_ventas = Detalleventa.objects.all()
    productos_vendidos = []

    for detalle in detalle_ventas:
        if len(productos_vendidos) == 0:
            productos_vendidos.append({"nombre": str(detalle.productoid), "cantidad":int(detalle.cantidad)})
        else:
            producto_found = next((product for product in productos_vendidos if product["nombre"] == str(detalle.productoid)), None)
            if producto_found:
                producto_found["cantidad"] = producto_found["cantidad"] +int(detalle.cantidad)
            else:
                productos_vendidos.append({"nombre": str(detalle.productoid), "cantidad":int(detalle.cantidad)})

    nombre_productos = []
    cantidad_productos = []

    for producto in productos_vendidos:
        nombre_productos.append(producto["nombre"]) 
        cantidad_productos.append(producto["cantidad"]) 

    productos_vendidos = []
    productos_vendidos.append(nombre_productos)
    productos_vendidos.append(cantidad_productos)
    productos_vendidos = json.dumps(productos_vendidos)

    productos_stock = Producto.objects.all().order_by('stock')[:10]
    productos_peligro_stock = [] 
    nombre = []
    stock = []

    for pro_stock in productos_stock:
        nombre.append(str(pro_stock.nombre))
        stock.append(int(pro_stock.stock)-int(pro_stock.stockcritico))

    productos_peligro_stock.append(nombre)
    productos_peligro_stock.append(stock)
    productos_peligro_stock = json.dumps(productos_peligro_stock)
    
    context = {
        'ventasXDocumento': ventas_tipodocumento,
        'ventasXDespacho': total_ventas_despachos,
        'productosXcantidad': productos_vendidos,
        'totalVentas':len(ventas),
        'productoStock':productos_peligro_stock,
        'tipo_usuario': tipo_usuario,

    }

    return render(request, 'dashboard.html', context)
